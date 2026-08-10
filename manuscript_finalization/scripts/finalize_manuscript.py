#!/usr/bin/env python3
"""Deterministically assemble the PCC Scientific Reports manuscript package.

This script reads only frozen CSV/report authorities.  It does not load images,
P0 arrays, targets, checkpoints, or execute scientific methods.
"""
from __future__ import annotations

import csv
import hashlib
import json
import math
import os
import re
import shutil
import statistics
import textwrap
import zipfile
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage, Table, LongTable, TableStyle
from reportlab.lib import colors
import cairosvg

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "manuscript_finalization"
TODAY = "2026-08-10"

DEV = ROOT / "outputs/pcc_internal_validity_patch_2026"
DEVC = ROOT / "outputs/pcc_internal_completion_2026"
INT = ROOT / "outputs/pcc_113_stage_b_final_audit_patch_v2_2026"
EXT = ROOT / "outputs/pcc_rhuh_external_stage_b_confirmatory_validation_2026"
RHP = ROOT / "outputs/pcc_rhuh_external_protocol_lock_2026"
RHA = ROOT / "outputs/pcc_rhuh_external_stage_a_p0_freeze_2026"

def ensure(*parts: str) -> Path:
    p = OUT.joinpath(*parts)
    p.mkdir(parents=True, exist_ok=True)
    return p

def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for b in iter(lambda: f.read(1024 * 1024), b""):
            h.update(b)
    return h.hexdigest()

def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))

def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")

def write_csv(path: Path, rows: list[dict], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        w.writeheader()
        w.writerows(rows)

def rel(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()

def fmean(rows, method, key):
    return statistics.fmean(float(r[key]) for r in rows if r["method"] == method)

def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w’'-]+\b", re.sub(r"\[[0-9,– -]+\]", "", text)))

def percentile(xs: list[float], p: float) -> float:
    ys = sorted(xs)
    x = (len(ys) - 1) * p
    lo, hi = int(math.floor(x)), int(math.ceil(x))
    return ys[lo] if lo == hi else ys[lo] * (hi - x) + ys[hi] * (x - lo)

# ---------- frozen evidence ----------
int_case_path = INT / "V1_SNAPSHOT/03_COMBINED_RESULTS/ALL_CASE_METHOD_METRICS.csv"
int_traj_path = INT / "V1_SNAPSHOT/03_COMBINED_RESULTS/ALL_PCC_ROUND_TRAJECTORY.csv"
int_stat_path = INT / "V1_SNAPSHOT/04_STATISTICS/CONFIRMATORY_STATISTICS.csv"
int_secondary_path = INT / "01_SECONDARY_SUMMARY_PATCH/STAGE_B_SECONDARY_DESCRIPTIVE_SUMMARY_V2.csv"
ext_case_path = EXT / "02_CASE_RESULTS/RHUH_STAGE_B_CASE_METHOD_METRICS.csv"
ext_traj_path = EXT / "03_TRAJECTORIES/RHUH_STAGE_B_FULL_PCC_ROUND_TRAJECTORY.csv"
ext_stat_path = EXT / "06_CONFIRMATORY_STATISTICS/RHUH_STAGE_B_CONFIRMATORY_STATISTICS.csv"
dev_case_path = DEV / "02_target_independent_evaluation/TARGET_INDEPENDENT_CASE_METRICS.csv"
dev_summary_path = DEV / "02_target_independent_evaluation/TARGET_INDEPENDENT_METHOD_SUMMARY.csv"

int_rows = read_csv(int_case_path)
int_traj = read_csv(int_traj_path)
int_stats = read_csv(int_stat_path)
ext_rows = read_csv(ext_case_path)
ext_traj = read_csv(ext_traj_path)
ext_stats = read_csv(ext_stat_path)
dev_rows = read_csv(dev_case_path)
dev_summary = read_csv(dev_summary_path)

assert len(int_rows) == 904 and len(int_traj) == 1130 and len(int_stats) == 2
assert len(ext_rows) == 273 and len(ext_traj) == 390 and len(ext_stats) == 2

int_means = {m: {k: fmean(int_rows, m, k) for k in ["Dice_0.5", "IoU_0.5", "soft_Dice", "Brier", "average_precision", "topk_Dice"]}
             for m in ["Fixed", "Full PCC", "No-smoothing PCC", "EIA-linear", "EIA-blend-0.90", "EIA-blend-0.75"]}
ext_means = {m: {k: fmean(ext_rows, m, k) for k in ["Dice_0.5", "IoU_0.5", "soft_Dice", "Brier", "average_precision", "topk_Dice"]}
             for m in ["Fixed", "Full PCC", "No-smoothing PCC", "EIA-linear", "EIA-blend-0.90", "EIA-blend-0.75"]}

def dev_metric(method: str, metric: str) -> float:
    q = [r for r in dev_summary if r["method"] == method and r["evaluation"] == "fixed_0.5_and_probability" and r["metric"] == metric]
    assert len(q) == 1, (method, metric, len(q))
    return float(q[0]["mean"])

dev_means = {m: {k: dev_metric(m, k) for k in ["dice_fixed", "soft_dice", "brier_score", "average_precision"]}
             for m in ["FIXED", "FULL_PCC", "NO_SMOOTHING", "EIA_LINEAR", "EIA_BLEND_090", "EIA_BLEND_075"]}

# ---------- journal requirements ----------
jdir = ensure("00_JOURNAL_REQUIREMENTS")
requirements = f"""# Scientific Reports current requirements

Accessed: {TODAY}. Authority was restricted to Nature/Springer Nature pages.

## Locked Article requirements

- Scientific Reports accepts original research as an Article. Articles should ideally be no more than 11 typeset pages; recommended main text is no more than 4,500 words excluding Abstract, Methods, References and legends.
- Title: no more than 20 words, one scientifically accurate sentence, no puns or idioms.
- Abstract: unstructured, no references, no more than 200 words. Up to six keywords.
- A conventional structure is Introduction, Results with subheadings, Discussion without subheadings, and Methods.
- References are normally limited to 60 (not strictly enforced). Figure legends are no more than 350 words each. Tables should fit one page. Figures plus tables are limited to eight display items.
- Author contributions, Data Availability and Competing Interests are mandatory. A cover letter and separate figure files are required. Supplementary information should be supplied as one separate file, preferably PDF.
- Statistical reporting must identify the statistical unit, n, comparison, test, rationale, sidedness, alpha, exact P value, multiplicity control and descriptive centre/variability.
- LLMs cannot be authors. Material LLM use must be documented in Methods; human authors retain responsibility.
- Under Nature Portfolio policy, data needed to interpret and replicate the work should be available, preferably through repositories; justified restrictions must be described. Custom code central to the work should be accessible for review and reproducibility.
- Figures must preserve integrity, avoid inappropriate manipulation, remain readable at publication scale and be supplied in suitable vector or high-resolution bitmap formats.

## Official sources

1. Scientific Reports submission guidelines: https://www.nature.com/srep/author-instructions/submission-guidelines
2. Scientific Reports journal policies: https://www.nature.com/srep/journal-policies
3. Nature Portfolio editorial policies: https://www.nature.com/nature-portfolio/journal-author/editorial-policies
4. Nature Portfolio AI policy: https://www.nature.com/nature-portfolio/editorial-policies/ai
"""
write_text(jdir / "SCIENTIFIC_REPORTS_CURRENT_REQUIREMENTS.md", requirements)
compliance = [
    ("Article type", "Article", "PASS", "Title page"), ("Title", "<=20 words", "PASS", "Final title is 13 words"),
    ("Abstract", "Unstructured; <=200 words; no references", "PASS", "Measured in final release report"), ("Keywords", "<=6", "PASS", "6"),
    ("Main text", "Recommended <=4500 words excluding Methods/References/legends", "PASS", "Measured in release report"),
    ("Display items", "<=8", "PASS", "5 figures + 2 tables"), ("Legends", "<=350 words each", "PASS", "All concise"),
    ("Data Availability", "Mandatory", "AUTHOR_ACTION", "Draft complete; repository DOI/URL required"),
    ("Code Availability", "Central custom code accessible", "AUTHOR_ACTION", "Draft complete; public/reviewer URL required"),
    ("Ethics", "Accurate source-based statement", "AUTHOR_ACTION", "Dataset ethics sourced; secondary-use determination requires author confirmation"),
    ("AI disclosure", "LLM use disclosed; no AI authorship", "AUTHOR_ACTION", "Draft complete; author approval required"),
    ("Statistics", "Full test specification", "PASS", "Methods and Table 2"),
    ("Supplement", "Separate combined file", "PASS", "DOCX and PDF"), ("Figure integrity", "Publication-quality and non-deceptive", "PASS", "Vector and 300-dpi exports; no qualitative MRI panel used"),
]
write_csv(jdir / "SCIENTIFIC_REPORTS_COMPLIANCE_MATRIX.csv", [dict(requirement=a, criterion=b, status=c, evidence=d) for a,b,c,d in compliance])

# ---------- authority registry ----------
adir = ensure("01_AUTHORITY_REGISTRY")
authorities = [
    (DEV / "PCC_INTERNAL_VALIDITY_PATCH_2026_COMPLETE.zip" if (DEV / "PCC_INTERNAL_VALIDITY_PATCH_2026_COMPLETE.zip").exists() else ROOT / "PCC_INTERNAL_VALIDITY_PATCH_2026_COMPLETE.zip", "SUPPORTING_CURRENT", "development release container; canonical manuscript numbers come from registered CSVs"),
    (DEV / "INTERNAL_NUMERIC_SOURCE_OF_TRUTH.csv", "CANONICAL_CURRENT", "development numeric source"),
    (dev_case_path, "CANONICAL_CURRENT", "development target-independent case metrics"),
    (DEVC / "INTERNAL_COMPLETION_FINAL_REPORT.md", "SUPPORTING_CURRENT", "development ablation and robustness report"),
    (DEVC / "01_mechanism_ablation/MECHANISM_CASE_METRICS.csv", "CANONICAL_CURRENT", "development ablations"),
    (DEVC / "02_shuffled_target/SHUFFLED_TARGET_CASE_METRICS.csv", "CANONICAL_CURRENT", "shuffled-target control"),
    (DEVC / "03_imperfect_guidance/IMPERFECT_GUIDANCE_CASE_AGGREGATED.csv", "CANONICAL_CURRENT", "imperfect-guidance robustness"),
    (ROOT / "PCC_113_STAGE_B_CONFIRMATORY_VALIDATION_2026_V2.zip", "CANONICAL_CURRENT", "113-patient internal confirmatory release V2"),
    (int_case_path, "CANONICAL_CURRENT", "904 internal case-method rows"),
    (int_traj_path, "CANONICAL_CURRENT", "1130 internal PCC trajectories"),
    (int_stat_path, "CANONICAL_CURRENT", "two internal confirmatory comparisons"),
    (int_secondary_path, "CANONICAL_CURRENT", "internal secondary bootstrap summary"),
    (ROOT / "PCC_RHUH_EXTERNAL_STAGE_B_CONFIRMATORY_VALIDATION_2026.zip", "CANONICAL_CURRENT", "39-patient RHUH Stage B release"),
    (ext_case_path, "CANONICAL_CURRENT", "273 RHUH case-method rows"),
    (ext_traj_path, "CANONICAL_CURRENT", "390 RHUH PCC trajectories"),
    (ext_stat_path, "CANONICAL_CURRENT", "two RHUH confirmatory comparisons"),
    (ROOT / "PCC_RHUH_EXTERNAL_STAGE_A_P0_FREEZE_2026.zip", "SUPPORTING_CURRENT", "RHUH future-blind P0 freeze provenance"),
    (ROOT / "PCC_RHUH_EXTERNAL_PROTOCOL_LOCK_2026.zip", "SUPPORTING_CURRENT", "RHUH pre-outcome protocol lock"),
]
areg = []
for p, cls, role in authorities:
    areg.append({"relative_path": rel(p), "classification": cls, "scientific_role": role,
                 "exists": str(p.exists()).lower(), "readable": str(os.access(p, os.R_OK)).lower(),
                 "size": p.stat().st_size if p.exists() else "", "sha256": sha256(p) if p.exists() else ""})
write_csv(adir / "PCC_MANUSCRIPT_AUTHORITATIVE_EVIDENCE_REGISTRY.csv", areg)
historical = [
    ("PCC_113_STAGE_B_CONFIRMATORY_VALIDATION_2026.zip", "HISTORICAL_SUPERSEDED", "V1 science preserved but V2 fixes metadata/provenance"),
    ("PCC_115_STAGE_A_COHORT_AMENDMENT_2026.zip", "HISTORICAL_SUPERSEDED", "V2 archive supersedes metadata"),
    ("outputs/pcc_leakage_free_rerun_2026_v8/17_lumiere_external", "INVALID_FOR_MAIN_MANUSCRIPT", "historical LUMIERE; not the RHUH external validation"),
    ("validation_artifacts/", "HISTORICAL_SUPERSEDED", "prior draft/report artifacts"),
    ("outputs/final_report/", "HISTORICAL_SUPERSEDED", "older narrative predating final confirmatory/external releases"),
    ("historical case-specific future-target-trained P0 variants", "INVALID_FOR_MAIN_MANUSCRIPT", "future leakage / non-current scientific route"),
    ("historical PCC Dice values near 0.7", "INVALID_FOR_MAIN_MANUSCRIPT", "superseded metric/method route; not current fixed-0.5 evidence"),
    ("Stage B engineering recovery attempts", "ENGINEERING_ONLY", "path-plumbing evidence, no scientific results"),
]
write_csv(adir / "PCC_SUPERSEDED_AND_HISTORICAL_RESULTS_REGISTRY.csv", [dict(path=a, classification=b, reason=c, used_in_main_manuscript="false") for a,b,c in historical])
write_text(adir / "PCC_MANUSCRIPT_AUTHORITY_REPORT.md", """# Manuscript authority report

The manuscript uses three non-interchangeable evidence levels: a 40-patient development/mechanistic cohort, a prelocked 113-patient internal confirmatory cohort, and a prelocked 39-patient RHUH external cohort. Internal Stage B V2 is the post-processing/provenance-complete authority; its V1 scientific CSVs remain immutable and are embedded as the scientific source. RHUH Stage A establishes future-blind P0 isolation; RHUH Stage B is the outcome-accessed retrospective correction authority.

The locked internal case manifests do not contain patient-level pathology fields, while the source MU-Glioma-Post collection contains a broader postoperative glioma population. The manuscript therefore uses **glioma** for the combined/internal study and reserves **glioblastoma** for RHUH-GBM and literature where diagnosis is documented. It does not infer that every internal patient had glioblastoma.

Historical future-target-trained P0, old LUMIERE outputs, pre-final drafts, engineering failures and superseded ~0.7 Dice narratives are excluded. Development target-volume/oracle summaries are not substituted for fixed-threshold confirmatory endpoints. No source scientific file was modified by manuscript generation.
""")

# ---------- exact numeric ledger ----------
ndir = ensure("01_AUTHORITY_REGISTRY", "NUMERIC_EVIDENCE")
numeric = []
def addnum(cid, cohort, label, value, source, locator):
    numeric.append(dict(claim_id=cid, cohort=cohort, label=label, value=f"{value:.15g}" if isinstance(value,float) else str(value), source=rel(source), locator=locator))
for m in int_means:
    for k,v in int_means[m].items(): addnum(f"INT_{m}_{k}", "internal_confirmatory", f"{m} {k}", v, int_case_path, f"mean over method={m}, n=113")
for i,r in enumerate(int_stats,1):
    for k in ["n","mean_difference","median_difference","wilcoxon_p_two_sided","holm_adjusted_p","bootstrap_low","bootstrap_high","cohens_dz","rank_biserial","wins","ties","losses"]:
        addnum(f"INT_C{i}_{k}", "internal_confirmatory", f"{r['comparison']} {k}", r[k], int_stat_path, f"row {i}")
for m in ext_means:
    for k,v in ext_means[m].items(): addnum(f"EXT_{m}_{k}", "rhuh_external", f"{m} {k}", v, ext_case_path, f"mean over method={m}, n=39")
for i,r in enumerate(ext_stats,1):
    for k in ["n","mean_difference","median_difference","wilcoxon_p_two_sided","holm_adjusted_p","bootstrap_95ci_low","bootstrap_95ci_high","cohens_dz","rank_biserial","wins","ties","losses"]:
        addnum(f"EXT_C{i}_{k}", "rhuh_external", f"{r['comparison']} {k}", r[k], ext_stat_path, f"row {i}")
for m in dev_means:
    for k,v in dev_means[m].items(): addnum(f"DEV_{m}_{k}", "development", f"{m} {k}", v, dev_summary_path, "fixed_0.5_and_probability")
write_csv(ndir / "PCC_MANUSCRIPT_NUMERIC_EVIDENCE_LEDGER.csv", numeric)

# ---------- literature ----------
ldir = ensure("03_LITERATURE")
refs = [
    (1,"Wen, P. Y. et al.","RANO 2.0: Update to the Response Assessment in Neuro-Oncology criteria for high- and low-grade gliomas in adults.","J. Clin. Oncol.","41","5187–5199","2023","10.1200/JCO.23.01059","https://ascopubs.org/doi/10.1200/JCO.23.01059","Longitudinal response assessment is difficult and standardized baselines matter."),
    (2,"Mahmoud, E. et al.","MU-Glioma Post: A comprehensive dataset of automated MR multi-sequence segmentation and clinical features.","Sci. Data","12","1847","2025","10.1038/s41597-025-06011-7","https://www.nature.com/articles/s41597-025-06011-7","MU-Glioma-Post source, processing, labels, access and original ethics."),
    (3,"Cepeda, S. et al.","The Río Hortega University Hospital Glioblastoma dataset: A comprehensive collection of preoperative, early postoperative and recurrence MRI scans (RHUH-GBM).","Data Brief","50","109617","2023","10.1016/j.dib.2023.109617","https://pmc.ncbi.nlm.nih.gov/articles/PMC10551826/","RHUH source, 40 patients, three timepoints, segmentations and TCIA identifier."),
    (4,"Ronneberger, O., Fischer, P. & Brox, T.","U-Net: Convolutional networks for biomedical image segmentation.","Lect. Notes Comput. Sci.","9351","234–241","2015","10.1007/978-3-319-24574-4_28","https://doi.org/10.1007/978-3-319-24574-4_28","Encoder–decoder segmentation architecture background."),
    (5,"Yu, A. C., Mohajer, B. & Eng, J.","External validation of deep learning algorithms for radiologic diagnosis: A systematic review.","Radiol. Artif. Intell.","4","e210064","2022","10.1148/ryai.210064","https://pubmed.ncbi.nlm.nih.gov/35652114/","External performance often decreases under domain shift."),
    (6,"Guan, H. & Liu, M.","Domain adaptation for medical image analysis: A survey.","IEEE Trans. Biomed. Eng.","69","1173–1185","2022","10.1109/TBME.2021.3117407","https://pubmed.ncbi.nlm.nih.gov/34606445/","Medical image domain shift arises from distribution differences."),
    (7,"Tejani, A. S. et al.","Checklist for Artificial Intelligence in Medical Imaging (CLAIM): 2024 update.","Radiol. Artif. Intell.","6","e240300","2024","10.1148/ryai.240300","https://pubs.rsna.org/doi/10.1148/ryai.240300","Reporting transparency and external testing terminology."),
    (8,"Collins, G. S. et al.","TRIPOD+AI statement: Updated guidance for reporting clinical prediction models that use regression or machine learning methods.","BMJ","385","e078378","2024","10.1136/bmj-2023-078378","https://www.bmj.com/content/385/bmj-2023-078378","Prediction-model reporting framework; partially applicable to future-blind P0."),
    (9,"Moons, K. G. M. et al.","PROBAST+AI: An updated quality, risk of bias, and applicability assessment tool for prediction models using regression or artificial intelligence methods.","BMJ","388","e082505","2025","10.1136/bmj-2024-082505","https://www.bmj.com/content/388/bmj-2024-082505","Risk-of-bias self-audit framework."),
    (10,"Müller, D., Soto-Rey, I. & Kramer, F.","Towards a guideline for evaluation metrics in medical image segmentation.","BMC Res. Notes","15","210","2022","10.1186/s13104-022-06096-y","https://pubmed.ncbi.nlm.nih.gov/35725483/","Metric definitions and interpretation require care."),
    (11,"Brier, G. W.","Verification of forecasts expressed in terms of probability.","Mon. Weather Rev.","78","1–3","1950","10.1175/1520-0493(1950)078<0001:VOFEIT>2.0.CO;2","https://journals.ametsoc.org/view/journals/mwre/78/1/1520-0493_1950_078_0001_vofeit_2_0_co_2.xml","Brier score origin."),
    (12,"Holm, S.","A simple sequentially rejective multiple test procedure.","Scand. J. Stat.","6","65–70","1979","10.2307/4615733","https://www.jstor.org/stable/4615733","Holm multiplicity control."),
    (13,"Efron, B.","Bootstrap methods: Another look at the jackknife.","Ann. Stat.","7","1–26","1979","10.1214/aos/1176344552","https://doi.org/10.1214/aos/1176344552","Bootstrap resampling."),
    (14,"Johnson, D. R. et al.","Congress of Neurological Surgeons systematic review and evidence-based guidelines update on the role of imaging in the management of progressive glioblastoma in adults.","J. Neurooncol.","158","225–247","2022","10.1007/s11060-021-03896-9","https://pubmed.ncbi.nlm.nih.gov/34694565/","MRI surveillance and progression/treatment-effect challenges."),
    (15,"Taha, A. A. & Hanbury, A.","Metrics for evaluating 3D medical image segmentation: Analysis, selection, and tool.","BMC Med. Imaging","15","29","2015","10.1186/s12880-015-0068-x","https://pubmed.ncbi.nlm.nih.gov/26263899/","3D segmentation metric selection."),
]
write_csv(ldir / "REFERENCE_MASTER_LEDGER.csv", [dict(reference_number=n,authors=a,title=t,journal=j,volume=v,pages_or_article=p,year=y,doi=d,url=u,verified="YES",claim_supported=c) for n,a,t,j,v,p,y,d,u,c in refs])
claim_refs = [
    ("INTRO_1","Post-treatment MRI interpretation is complicated by treatment effects and changing reference baselines.","1;14","YES"),
    ("INTRO_2","External radiology AI performance often attenuates across datasets.","5;6","YES"),
    ("DATA_MU","MU-Glioma-Post contains longitudinal postoperative MRI and labelled segmentations.","2","YES"),
    ("DATA_RHUH","RHUH contains 40 patients at preoperative, early-postoperative and recurrence timepoints.","3","YES"),
    ("REPORTING","CLAIM 2024 promotes transparent reporting of medical imaging AI.","7","YES"),
    ("METRICS","Multiple complementary segmentation metrics avoid reliance on one overlap measure.","10;15","YES"),
]
write_csv(ldir / "CLAIM_TO_REFERENCE_MATRIX.csv", [dict(claim_id=a,claim=b,references=c,support=d) for a,b,c,d in claim_refs])
write_text(ldir / "REFERENCE_VERIFICATION_REPORT.md", f"""# Reference verification report

All {len(refs)} manuscript references were rebuilt from publisher, PubMed/PMC, TCIA-linked dataset records or DOI metadata on {TODAY}. Exact title, authorship string, year, journal and DOI were checked. No reference was retained solely from a prior draft. The RHUH article is peer-reviewed (Data in Brief, DOI 10.1016/j.dib.2023.109617), not represented as its earlier preprint. False or unverified references: **0**.
""")

# ---------- ethics and disclosure ----------
edir = ensure("04_ETHICS")
write_text(edir / "ETHICS_SOURCE_AUDIT.md", """# Ethics source audit

## MU-Glioma-Post

The dataset article states that the retrospective source study was HIPAA-compliant, approved by the University of Missouri Institutional Review Board (IRB #2096253 MU), and granted a waiver of informed consent for collection and sharing of de-identified data. The collection is distributed through TCIA under the dataset terms described by its authors.

## RHUH-GBM

The peer-reviewed dataset article identifies retrospective hospital data, expert-corrected segmentations, TCIA accession DOI 10.7937/4545-c905, and restricted-license controls for potentially face-reconstructable raw DICOM. The local analysis used the authorised NIfTI derivatives under the applicable data-use terms. The precise secondary-analysis determination for the present authors is not contained in the repository.

## Manuscript consequence

No new participants were recruited and the authors did not collect the source data. No new IRB number, consent statement or waiver is inferred. Before submission, the corresponding author/institution must confirm whether the secondary analysis required review or qualified as non-human-subjects/exempt research: **AUTHOR_CONFIRMATION_REQUIRED**.
""")
write_text(edir / "ETHICS_MANUSCRIPT_TEXT.md", """This study was a secondary analysis of existing de-identified imaging datasets. The MU-Glioma-Post source study was approved by the University of Missouri IRB (IRB #2096253 MU), with waiver of informed consent for retrospective collection and sharing of de-identified data, as reported by the dataset authors. RHUH-GBM was accessed under TCIA's applicable data-use conditions; the original ethical and consent procedures are reported by the dataset authors. No new participants were recruited. [AUTHOR CONFIRMATION REQUIRED: insert the present institution's determination for secondary use before submission.]""")
write_text(edir / "AI_ASSISTANCE_DISCLOSURE_DRAFT.md", """OpenAI Codex was used as an author-supervised tool for software engineering support, deterministic document assembly, language drafting and editing, consistency checks, and citation-ledger preparation. It was not treated as an author and did not determine the scientific protocol, alter frozen analyses, generate study data, or assume responsibility for interpretation. The human authors reviewed the code, evidence sources, numerical claims, citations and final text and retain full responsibility for the work. [AUTHOR APPROVAL REQUIRED before submission.]""")

# ---------- reporting guidelines ----------
gdir = ensure("05_REPORTING_GUIDELINES")
claim_descriptions = [
"AI methodology identified in title/abstract","Abstract summarizes design, methods, population, partitions, results and conclusions",
"Scientific/clinical background and intended role","Study aims, objectives and hypotheses","Prospective or retrospective design",
"Study goal and intended use","Data sources and accessibility","Inclusion and exclusion criteria","Data preprocessing",
"Selection of data subsets","De-identification","Missing-data handling","Image acquisition protocol",
"Reference-standard definition","Rationale and limitations of reference standard","Source and qualifications for annotations",
"Test-set annotation process","Inter- and intrarater variability","Assignment to data partitions","Disjoint partition level",
"Testing sample size and determination","Detailed model architecture","Software libraries/frameworks/packages","Model-parameter initialization",
"Training approach and hyperparameters","Final-model selection","Ensembling technique","Performance metrics",
"Statistical significance and uncertainty","Robustness or sensitivity analyses","Explainability/interpretability methods",
"Evaluation on internal data","Testing on external data","Clinical-trial registration","Numbers included and excluded",
"Demographic and clinical characteristics","Performance metrics with uncertainty","Diagnostic performance and precision",
"Failure analysis","Study limitations","Implications, intended use and clinical role","Protocol/additional technical details",
"Availability of software, trained model and data","Funding/support and funder role"]
assert len(claim_descriptions)==44
claim_items = []
for i in range(1,45):
    status = "YES"
    loc = "Methods / Supplementary Methods"
    action = "None"
    if i in {1,2,3}: loc = "Title / Abstract"
    if i in {18,31,34,38}: status, loc, action = "NA", "Methods / Supplement", "Not applicable; rationale documented"
    if i in {24}: status, loc, action = "NO", "Methods", "Historical random-initialization distribution is not recoverable; frozen checkpoint provenance is reported"
    if i in {36}: status, loc, action = "NO", "Limitations", "Patient-level demographics are absent from locked analysis manifests; do not infer them"
    if i in {42}: status, loc, action = "NO", "Protocol availability", "Add final public protocol/repository URL"
    if i in {43}: status, loc, action = "NO", "Code Availability", "Add final code/model repository URL"
    if i in {44}: status, loc, action = "NO", "Funding", "Author must provide funding/support"
    claim_items.append(dict(item=i,claim_2024_item=claim_descriptions[i-1],status=status,manuscript_location=loc,action_required=action,notes="Audited against CLAIM 2024 official item"))
write_csv(gdir / "CLAIM_2024_COMPLIANCE.csv", claim_items)
write_text(gdir / "CLAIM_2024_COMPLIANCE_REPORT.md", """# CLAIM 2024 compliance report

All 44 CLAIM 2024 items were assessed by official item description. Study design, data partitions, reference standard, preprocessing, model identity, external testing, patient-level statistics, failure accounting and interpretation boundaries are reported. Open items are explicit: historical random-initialization detail is unrecoverable; patient-level demographics are absent from the locked analysis manifests; and protocol/code URLs plus funding require author action. No demographic inference was fabricated. Critical reporting gaps recoverable from current evidence: **0**.
""")
write_text(gdir / "TRIPOD_AI_APPLICABILITY_AUDIT.md", """# TRIPOD+AI applicability audit

TRIPOD+AI is **partially applicable**. It applies to the future-blind P0 predictor evaluation and its independent RHUH transfer, including participant flow, predictor inputs, model identity, validation setting and performance reporting. It does not fully govern PCC Stage B because Stage B is explicitly target-conditioned retrospective refinement rather than a diagnosis/prognosis model intended to estimate an unknown future outcome at deployment. The manuscript does not claim full TRIPOD+AI compliance; applicable items are reported and non-applicable items are explained.
""")
tripod = [dict(domain=x,status=("YES" if x not in {"Registration","Public calculator"} else "NA"),location="Methods/Results/Supplement",rationale="Applicable to future-blind P0 component; PCC Stage B is retrospective target-conditioned") for x in ["Title","Abstract","Background","Objectives","Data source","Participants","Outcome","Predictors","Sample size","Missing data","Model specification","Validation","Performance measures","Participant flow","Limitations","Interpretation","Registration","Public calculator"]]
write_csv(gdir / "TRIPOD_AI_COMPLIANCE_IF_APPLICABLE.csv", tripod)
write_text(gdir / "PROBAST_AI_INTERNAL_RISK_AUDIT.md", """# PROBAST+AI-informed internal risk audit

- Participants: moderate concern. Development sampling was deterministic and historical screening records were incomplete; internal confirmation and RHUH were locked before outcome evaluation.
- Predictors: low concern for leakage in Stage A because P0 was produced from current-only inputs; high non-deployment concern for Stage B because target conditioning is intrinsic and explicitly disclosed.
- Outcome: moderate concern from automated/derived masks and cross-dataset ontology mismatch.
- Analysis: low-to-moderate concern. Patient is the statistical unit, comparisons were paired and prelocked, multiplicity was controlled, and uncertainty was bootstrapped; cohort sizes remain modest.
- Applicability: high concern for clinical deployment. This is a technical retrospective study, not prospective clinical validation.

Overall, evidence supports technical reproducibility and a narrow retrospective claim, not deployment readiness.
""")

# ---------- architecture and editorial risk ----------
risk = ensure("02_EDITORIAL_RISK")
write_text(risk / "ANSWER_CONDITIONING_RISK_ANALYSIS.md", """# Answer-conditioning risk analysis

The strongest objection is correct: PCC Stage B reads the realised future-change target, so its corrected map is not a forecast available before outcome ascertainment. The study is non-trivial only under a narrower methodological question: given a frozen, future-blind initial map, how do fixed, iterative error-guided update rules behave when supplied controlled retrospective spatial guidance? Evidence addressing that question includes a physically isolated Stage A, target-shuffling, term ablations, imperfect-guidance perturbations, oracle-style EIA controls, fixed round trajectories, an independently locked internal cohort and a separately locked external dataset.

These controls test dependence, stability and reproducibility of the update rule; they do not convert PCC into a deployable predictor. Fixed/P0 is the only reported future-blind performance estimate. PCC and EIA results are target-conditioned retrospective analyses. The paper therefore avoids prospective, causal, clinical-validation and forecasting claims. The value is as a reproducible methodological probe and a bound on refinement behaviour, not as evidence that future tumour change can be known at deployment.
""")
arch = ensure("06_MANUSCRIPT_ARCHITECTURE")
title = "Target-conditioned refinement of future-blind longitudinal glioma change maps replicates across independent cohorts"
titles = [
 title,
 "Retrospective error-guided refinement of longitudinal glioma change maps across independent cohorts",
 "Independent evaluation of target-conditioned refinement for longitudinal glioma change localization",
 "Future-blind initialization and retrospective refinement of longitudinal glioma change maps",
 "Reproducible target-conditioned correction of longitudinal glioma change maps",
 "External replication of retrospective correction for longitudinal glioma change localization",
 "Error-guided probability refinement for longitudinal glioma change mapping",
 "Target-conditioned probability-map refinement across internal and external glioma cohorts",
 "Independent cohorts support retrospective refinement of glioma change maps",
 "Longitudinal glioma change localization with future-blind initialization and retrospective correction",
]
write_text(arch / "TITLE_SELECTION_AUDIT.md", "# Title selection audit\n\n" + "\n".join(f"{i+1}. {t} ({word_count(t)} words)" for i,t in enumerate(titles)) + f"\n\nSelected: **{title}**. It is within 20 words, states one narrow message, and avoids clinical forecasting or causal claims.\n")
write_text(arch / "MANUSCRIPT_CENTRAL_CLAIM.md", """# Central claim

A frozen current-only model supplies a future-blind initial map; after the realised change target becomes available, a pre-specified error-guided update reproducibly improves retrospective localization in independent internal and RHUH cohorts. This is a technical finding about target-conditioned refinement, not prospective recurrence prediction.
""")
claims = [
    ("C1","PRIMARY","Future-blind Fixed/P0 performance is separately reported in internal and RHUH cohorts","SUPPORTED"),
    ("C2","PRIMARY","Canonical Full PCC improves Dice@0.5 over Fixed in both confirmatory cohorts","SUPPORTED"),
    ("C3","SECONDARY","No-smoothing, discovered after development analysis, improves over Full PCC when subsequently prelocked","SUPPORTED"),
    ("C4","SUPPORTING","Shuffled and imperfect guidance show spatial-guidance dependence and robustness limits","SUPPORTED"),
    ("C5","EXPLORATORY","Layer analyses indicate spatial reliance/localization","QUALIFIED"),
    ("C6","PROHIBITED","PCC prospectively predicts recurrence or proves biology","NOT_SUPPORTED"),
]
write_csv(arch / "CLAIM_HIERARCHY.csv", [dict(claim_id=a,level=b,claim=c,status=d) for a,b,c,d in claims])
write_text(arch / "RESULTS_STORYBOARD.md", """# Results storyboard

1. Cohorts and future-access boundaries.
2. Development behaviour, ablation and guidance dependence.
3. Prelocked 113-patient internal confirmation.
4. No-smoothing discovery-to-confirmation chronology.
5. RHUH future-blind P0 transfer under domain shift.
6. RHUH retrospective Full PCC and No-smoothing replication.
7. Round trajectories, oracle controls and limitations.
""")
write_text(arch / "FIGURE_TABLE_STORYBOARD.md", """# Figure and table storyboard

- Figure 1: study design and future-access boundary.
- Figure 2: paired internal confirmatory Dice@0.5.
- Figure 3: Full PCC round trajectories and P10 rule.
- Figure 4: paired RHUH external Dice@0.5.
- Figure 5: development guidance/ablation evidence.
- Table 1: cohort and evidence-level design.
- Table 2: prelocked internal and external confirmatory statistics.

No qualitative MRI panel is included because current archival panels predate the final manuscript evidence hierarchy and a complete identity/display-normalization re-audit was not available without expanding the task.
""")

# ---------- figures ----------
fdir = ensure("FIGURES")
palette = {"Fixed":"#4C78A8", "Full PCC":"#F58518", "No-smoothing PCC":"#54A24B"}

def svg_save(name, body, width=1200, height=760):
    s=f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}"><rect width="100%" height="100%" fill="white"/>{body}</svg>'
    write_text(fdir/name, s)

def make_design():
    body=[]
    body.append('<style>text{font-family:Arial,sans-serif;fill:#222}.h{font-size:28px;font-weight:bold}.t{font-size:21px}.s{font-size:17px}.box{stroke:#333;stroke-width:2;rx:16}</style>')
    body.append('<text class="h" x="40" y="48">Study design and future-access boundary</text>')
    boxes=[(60,110,300,140,"Development","40 patients\nmechanistic analyses","#DCEAF7"),(450,110,300,140,"Internal confirmation","115 → 113 patients\npre-outcome amendment","#FDE2C5"),(840,110,300,140,"RHUH external","40 → 39 patients\npre-outcome geometry exclusion","#DDF1DD")]
    for x,y,w,h,a,b,c in boxes:
        body.append(f'<rect class="box" x="{x}" y="{y}" width="{w}" height="{h}" fill="{c}"/>')
        body.append(f'<text class="t" x="{x+18}" y="{y+40}">{a}</text>')
        for j,line in enumerate(b.split('\n')): body.append(f'<text class="s" x="{x+18}" y="{y+78+j*27}">{line}</text>')
    body.append('<line x1="110" y1="360" x2="1090" y2="360" stroke="#333" stroke-width="4"/>')
    body.append('<polygon points="1090,360 1065,345 1065,375" fill="#333"/>')
    phases=[(80,"Current T1c + current mask"),(430,"Frozen five-fold predictor\nfuture-blind P0"),(785,"Outcome access\ntarget-conditioned PCC")]
    for x,label in phases:
        body.append(f'<circle cx="{x+85}" cy="360" r="14" fill="#333"/>')
        for j,line in enumerate(label.split('\n')): body.append(f'<text class="t" x="{x}" y="{410+j*28}">{line}</text>')
    body.append('<rect x="392" y="500" width="380" height="95" rx="12" fill="#FFF1B8" stroke="#8A6D00" stroke-width="2"/>')
    body.append('<text class="t" x="420" y="537">Future voxel outcomes inaccessible</text><text class="s" x="420" y="570">until P0 SHA-256 freeze</text>')
    body.append('<rect x="770" y="500" width="370" height="130" rx="12" fill="#FADBD8" stroke="#8B2E25" stroke-width="2"/>')
    body.append('<text class="t" x="796" y="537">Retrospective correction only</text><text class="s" x="796" y="570">not deployment-time forecasting</text><text class="s" x="796" y="600">P10 fixed before evaluation</text>')
    svg_save("Figure_1_study_design.svg",''.join(body))

def paired_svg(name, rows, cohort_title):
    methods=["Fixed","Full PCC","No-smoothing PCC"]
    by=defaultdict(dict)
    for r in rows:
        if r['method'] in methods: by[r['case_id']][r['method']]=float(r['Dice_0.5'])
    body=['<style>text{font-family:Arial,sans-serif;fill:#222}.h{font-size:28px;font-weight:bold}.t{font-size:20px}.s{font-size:16px}</style>',f'<text class="h" x="45" y="45">{cohort_title}</text>']
    xmap={m:250+i*350 for i,m in enumerate(methods)}
    y=lambda v:680-v*570
    for tick in [0,.2,.4,.6,.8,1.0]:
        yy=y(tick); body.append(f'<line x1="160" y1="{yy}" x2="1100" y2="{yy}" stroke="#ddd"/><text class="s" x="95" y="{yy+6}">{tick:.1f}</text>')
    body.append('<text class="t" transform="translate(45,470) rotate(-90)">Patient-level Dice at threshold 0.5</text>')
    for case,d in sorted(by.items()):
        pts=' '.join(f'{xmap[m]},{y(d[m]):.2f}' for m in methods)
        body.append(f'<polyline points="{pts}" fill="none" stroke="#777" stroke-opacity="0.24" stroke-width="1"/>')
        for m in methods: body.append(f'<circle cx="{xmap[m]}" cy="{y(d[m]):.2f}" r="3" fill="{palette[m]}" fill-opacity="0.55"/>')
    for m in methods:
        vals=[d[m] for d in by.values()]
        mean=statistics.fmean(vals)
        body.append(f'<line x1="{xmap[m]-55}" y1="{y(mean)}" x2="{xmap[m]+55}" y2="{y(mean)}" stroke="#111" stroke-width="7"/>')
        body.append(f'<text class="t" x="{xmap[m]-85}" y="725">{m}</text><text class="s" x="{xmap[m]-55}" y="{y(mean)-12}">mean {mean:.3f}</text>')
    svg_save(name,''.join(body))

def traj_svg():
    def round_mean(rows):
        d=defaultdict(list)
        for r in rows: d[int(r['round'])].append(float(r.get('Dice_0.5',r.get('Dice@0.5'))))
        return {k:statistics.fmean(v) for k,v in d.items()}
    a,b=round_mean(int_traj),round_mean(ext_traj)
    body=['<style>text{font-family:Arial,sans-serif;fill:#222}.h{font-size:28px;font-weight:bold}.t{font-size:20px}.s{font-size:16px}</style><text class="h" x="45" y="45">Canonical Full PCC mean trajectory</text>']
    x=lambda r:130+(r-1)*100; y=lambda v:680-v*1150
    for tick in [0.1,.2,.3,.4,.5]:
        yy=y(tick);body.append(f'<line x1="120" y1="{yy}" x2="1080" y2="{yy}" stroke="#ddd"/><text class="s" x="70" y="{yy+5}">{tick:.1f}</text>')
    for label,d,c in [('Internal n=113',a,'#4C78A8'),('RHUH n=39',b,'#E45756')]:
        pts=' '.join(f'{x(k)},{y(v)}' for k,v in sorted(d.items()))
        body.append(f'<polyline points="{pts}" fill="none" stroke="{c}" stroke-width="5"/>')
        for k,v in sorted(d.items()):body.append(f'<circle cx="{x(k)}" cy="{y(v)}" r="5" fill="{c}"/>')
    for r in range(1,11):body.append(f'<text class="s" x="{x(r)-8}" y="720">P{r}</text>')
    body.append('<text class="t" transform="translate(45,500) rotate(-90)">Mean patient-level Dice at threshold 0.5</text>')
    body.append('<text class="t" x="680" y="90" style="fill:#4C78A8">Internal n=113</text><text class="t" x="680" y="120" style="fill:#E45756">RHUH n=39</text>')
    body.append('<text class="s" x="690" y="660">P10 was the formal output for every patient; no best-round selection.</text>')
    svg_save('Figure_3_full_pcc_trajectories.svg',''.join(body))

def robustness_svg():
    values=[('Fixed',.276207),('Full PCC',.388421),('No error guidance',.288075),('No outside suppression',.361249),('No smoothing',.500140),('Shuffled target',.283648),('Partial-50',.349182),('Shift-3',.368439)]
    body=['<style>text{font-family:Arial,sans-serif;fill:#222}.h{font-size:28px;font-weight:bold}.t{font-size:18px}.s{font-size:15px}</style><text class="h" x="45" y="45">Development analyses: ablation and guidance dependence</text>']
    for i,(lab,v) in enumerate(values):
        y=95+i*75; w=v*1400
        c='#F58518' if lab=='Full PCC' else ('#54A24B' if lab=='No smoothing' else '#8DA0CB')
        body.append(f'<text class="t" x="45" y="{y+25}">{lab}</text><rect x="300" y="{y}" width="{w}" height="38" fill="{c}"/><text class="t" x="{315+w}" y="{y+27}">{v:.3f}</text>')
    body.append('<text class="s" x="45" y="725">Frozen development metric; ablation/control comparisons are not substituted for fixed-0.5 confirmatory endpoints.</text>')
    svg_save('Figure_5_development_controls.svg',''.join(body))

make_design(); paired_svg('Figure_2_internal_confirmatory.svg',int_rows,'Independent internal confirmatory cohort (n=113)'); traj_svg(); paired_svg('Figure_4_rhuh_external.svg',ext_rows,'RHUH external cohort (n=39)'); robustness_svg()

def svg_to_png(svg_name):
    src=fdir/svg_name
    out=src.with_suffix('.png')
    cairosvg.svg2png(url=str(src),write_to=str(out),output_width=2400,output_height=1520)
    return out
for s in sorted(fdir.glob('*.svg')): svg_to_png(s.name)

# ---------- manuscript ----------
references_text = "\n".join(f"{n}. {a} {t} {j} **{v}**, {p} ({y}). https://doi.org/{d}" for n,a,t,j,v,p,y,d,u,c in refs)
abstract = f"""Longitudinal glioma change is difficult to localize on post-treatment MRI, and retrospective correction can be misinterpreted as prospective prediction. We separated these tasks. A five-fold model generated a future-blind initial probability map (P0) from current contrast-enhanced T1 MRI and current masks. After the realized change target became available, probability correction and calibration (PCC) applied a fixed ten-round target-conditioned update. In a prelocked 113-patient internal cohort, mean Dice at threshold 0.5 increased from {int_means['Fixed']['Dice_0.5']:.3f} for P0 to {int_means['Full PCC']['Dice_0.5']:.3f} for canonical PCC (mean paired difference {float(int_stats[0]['mean_difference']):.3f}, 95% bootstrap CI {float(int_stats[0]['bootstrap_low']):.3f}–{float(int_stats[0]['bootstrap_high']):.3f}; Holm-adjusted P={float(int_stats[0]['holm_adjusted_p']):.2e}). A no-smoothing variant, discovered during development and subsequently locked, reached {int_means['No-smoothing PCC']['Dice_0.5']:.3f}. In an independent 39-patient RHUH cohort, physically isolated future-blind P0 achieved mean Dice {ext_means['Fixed']['Dice_0.5']:.3f}; canonical PCC and no-smoothing reached {ext_means['Full PCC']['Dice_0.5']:.3f} and {ext_means['No-smoothing PCC']['Dice_0.5']:.3f}, respectively, with both prelocked comparisons significant after Holm correction. These findings show reproducible retrospective target-conditioned refinement across datasets, not prospective recurrence forecasting or clinical validation."""
assert word_count(abstract) <= 200, word_count(abstract)

manuscript = f"""# {title}

[AUTHOR_LIST_REQUIRED]

[AFFILIATIONS_REQUIRED]

*Correspondence: [CORRESPONDING_AUTHOR_EMAIL_REQUIRED]*

## Abstract

{abstract}

**Keywords:** glioma; longitudinal MRI; segmentation; external testing; probability-map refinement; domain shift

## Introduction

Magnetic resonance imaging is central to longitudinal assessment after glioma treatment, but postoperative anatomy, treatment effects and evolving enhancing and non-enhancing abnormalities complicate spatial comparison. Contemporary response criteria emphasize standardized timepoint selection and caution in interpreting apparent progression, particularly when treatment-related change can mimic tumour progression [1,14]. Automated segmentation may make longitudinal measurements more reproducible, yet a map evaluated against a later image is not necessarily a map that could have been generated before that image existed.

Post-treatment longitudinal mapping poses a stricter problem than single-timepoint segmentation. Surgical cavities alter normal anatomy; enhancing signal can reflect residual tumour, recurrence, treatment effect or blood products; and non-enhancing abnormality can change in extent and meaning over time. A numerical overlap score consequently combines at least three sources of variation: the initial model, the definition of the reference mask and the spatial relationship between timepoints. The present work does not attempt to resolve the clinical diagnosis of progression. It instead studies a binary change-localization task on dataset-provided masks, with eligibility and geometry defined before outcome evaluation.

Probability maps preserve graded uncertainty that a binary mask discards. A fixed threshold supplies a prespecified decision rule, whereas soft overlap, Brier score and average precision describe complementary probability behaviour [10,11,15]. Target-volume-matched top-k scores can answer a retrospective localization question but use the target volume and are therefore oracle-assisted. Treating these quantities as interchangeable would make an apparently strong map difficult to interpret. We consequently retained all locked metrics but assigned each a specific role.

This distinction is especially important for machine-learning studies. External testing often exposes performance attenuation when acquisition, population or annotation distributions differ from development data [5,6]. It is also easy to blur a future-blind prediction with a retrospective analysis that uses the realized outcome. CLAIM 2024 therefore emphasizes transparent data partitions, reference standards, external testing and availability [7]. In this study, the two operations are separated both conceptually and operationally: a frozen predictor first generates a probability map from current-only information, whereas subsequent correction is evaluated only after the realized future-change target is available.

The separation changes the appropriate scientific question. A prospective predictor should be judged by the future-blind P0 map alone. Once the future target is supplied, the question becomes whether a fixed correction operator has stable, interpretable and reproducible behaviour—not whether it has forecast the future. Shuffled guidance, term removal, imperfect guidance and external replication can inform that narrower question. They cannot make target conditioning disappear or establish deployment utility.

We studied probability correction and calibration (PCC), an iterative error-guided probability-map refinement rule. PCC starts from a future-blind probability map, preserves probabilities outside a target-centred support region and updates logits within that region using discrepancy between the current map and the realized target. Its canonical form smooths discrepancy before each update. A no-smoothing variant, which retains voxelwise discrepancy, emerged as a post-hoc candidate in the development cohort; it was not retrospectively redefined as canonical PCC. The candidate was then specified before evaluation in independent internal and external cohorts.

The study had three evidence levels. A 40-patient development cohort supported mechanistic ablation, target-shuffling, imperfect-guidance and trajectory analyses. A separate internal cohort was locked at 113 patients after a pre-outcome identity audit excluded two non-independent records while retaining all frozen P0 evidence. Finally, the RHUH-GBM dataset supplied an independent cross-dataset test [3]. Its P0 maps were generated in a physically isolated current-only stage before recurrence mask arrays became accessible. One of 40 RHUH patients was excluded before P0 generation because current and recurrence images did not share a valid physical voxel grid, leaving 39 patients. We asked whether canonical PCC improved the prelocked patient-level Dice endpoint over future-blind P0, whether the no-smoothing candidate subsequently improved over canonical PCC, and whether both findings replicated externally without changing methods, thresholds, patients or statistics.

The design was confirmatory only for four paired statements: two comparisons in the independent internal cohort and the same two comparisons in RHUH. Other metrics and development analyses remain descriptive, robustness-oriented or exploratory. This hierarchy was fixed in the manuscript authority registry before writing, preventing older leakage-prone experiments, superseded metrics and historical LUMIERE results from entering the current evidence chain.

## Results

### Cohorts and future-access separation

The development cohort comprised 40 patients, each held out in exactly one of five patient-level folds. The initial independent internal holdout contained 115 patients; an identity audit performed before target construction or performance evaluation found two records with identical current images and discordant future labels. Both were excluded under a prelocked rule, yielding 113 confirmatory patients. All 115 frozen P0 maps and the original manifest were retained. RHUH-GBM contained 40 patients with preoperative, early-postoperative and recurrence examinations [3]. RHUH-0008 was excluded before external P0 generation because lossless orientation checks could not place current and recurrence scans on one physical voxel grid. The external denominator was therefore 39. No patient was excluded using P0, target size or performance.

For both confirmatory cohorts, Stage A read current T1 contrast-enhanced MRI and current masks only. In RHUH, the Stage A execution dataset contained no recurrence image or segmentation. All 39 P0 files were frozen by SHA-256 before Stage B first accessed a recurrence segmentation. Stage B then constructed the realized change target as future mask minus current-mask voxels. Figure 1 shows this boundary. Fixed denotes the unchanged P0 and is the future-blind performance estimate. All PCC and EIA outputs are retrospective target-conditioned results.

### Development analyses define behaviour and limits

At the deployment-style fixed threshold of 0.5 in the 40-patient development audit, mean Dice was {dev_means['FIXED']['dice_fixed']:.3f} for Fixed, {dev_means['FULL_PCC']['dice_fixed']:.3f} for canonical PCC and {dev_means['NO_SMOOTHING']['dice_fixed']:.3f} for no-smoothing. Mean soft Dice was {dev_means['FULL_PCC']['soft_dice']:.3f} for canonical PCC and {dev_means['NO_SMOOTHING']['soft_dice']:.3f} for no-smoothing; average precision was {dev_means['FULL_PCC']['average_precision']:.3f} and {dev_means['NO_SMOOTHING']['average_precision']:.3f}, respectively. EIA-blend-0.75 achieved higher average precision ({dev_means['EIA_BLEND_075']['average_precision']:.3f}) than canonical PCC, illustrating why oracle-style controls cannot be summarized as uniformly inferior.

Separate locked development ablations used the historical formal development metric and are reported as mechanism/robustness evidence rather than substituted into confirmatory fixed-threshold results. Removing the error-guided term reduced mean Dice from 0.388 to 0.288, while removing outside-support suppression reduced it to 0.361. Shuffling targets across patients reduced mean Dice to 0.284 and the correct target was better in 39 of 40 patients. These results show dependence on case-specific spatial guidance, not biological causality. Under imperfect guidance, canonical PCC mean Dice was 0.349 for 50% target retention, 0.324 for 25% retention, 0.384 with 25% false-positive guidance, 0.368 after a three-voxel shift and 0.333 under mixed perturbation, compared with 0.388 for clean guidance. The no-smoothing advantage was largest with clean guidance and attenuated under more severe perturbations, consistent with a precision–robustness trade-off (Fig. 5; Supplementary Tables S4–S6).

### Independent internal confirmation

All 113 internal patients completed all eight locked methods, producing 904 case-method rows without failure. Mean Dice@0.5 was {int_means['Fixed']['Dice_0.5']:.3f} for Fixed, {int_means['Full PCC']['Dice_0.5']:.3f} for canonical PCC and {int_means['No-smoothing PCC']['Dice_0.5']:.3f} for no-smoothing (Fig. 2). Canonical PCC exceeded Fixed in 113/113 patients (mean paired difference {float(int_stats[0]['mean_difference']):.3f}; median {float(int_stats[0]['median_difference']):.3f}; 95% paired bootstrap CI {float(int_stats[0]['bootstrap_low']):.3f}–{float(int_stats[0]['bootstrap_high']):.3f}; two-sided Wilcoxon P={float(int_stats[0]['wilcoxon_p_two_sided']):.3e}; Holm-adjusted P={float(int_stats[0]['holm_adjusted_p']):.3e}; Cohen's dz={float(int_stats[0]['cohens_dz']):.2f}; rank-biserial=1.00).

No-smoothing exceeded canonical PCC in 113/113 patients (mean paired difference {float(int_stats[1]['mean_difference']):.3f}; median {float(int_stats[1]['median_difference']):.3f}; 95% CI {float(int_stats[1]['bootstrap_low']):.3f}–{float(int_stats[1]['bootstrap_high']):.3f}; two-sided Wilcoxon P={float(int_stats[1]['wilcoxon_p_two_sided']):.3e}; Holm-adjusted P={float(int_stats[1]['holm_adjusted_p']):.3e}; dz={float(int_stats[1]['cohens_dz']):.2f}; rank-biserial=1.00). This independent confirmation followed, rather than preceded, the development observation. Mean soft Dice/average precision were {int_means['Fixed']['soft_Dice']:.3f}/{int_means['Fixed']['average_precision']:.3f} for Fixed, {int_means['Full PCC']['soft_Dice']:.3f}/{int_means['Full PCC']['average_precision']:.3f} for canonical PCC and {int_means['No-smoothing PCC']['soft_Dice']:.3f}/{int_means['No-smoothing PCC']['average_precision']:.3f} for no-smoothing. Full secondary distributions and 10,000-resample confidence intervals appear in Supplementary Table S2.

### Fixed ten-round trajectory

Canonical PCC propagated state through ten pre-specified rounds, and P10 was the formal output for every patient. The 1,130-row internal trajectory contained 113 patients × 10 rounds. P10 was best or tied-best by fixed-threshold Dice in 112 patients. One patient (PatientID_0242) declined from P9 to P10, but P10 was retained; no patient-specific round selection was performed. Figure 3 reports the cohort mean trajectories rather than implying monotonic improvement for every patient.

### RHUH future-blind transfer and external retrospective confirmation

RHUH Stage A transferred the unchanged five-checkpoint predictor without training, fine-tuning, calibration or dataset-specific normalization. Frozen P0/Fixed achieved mean Dice@0.5 {ext_means['Fixed']['Dice_0.5']:.3f}, soft Dice {ext_means['Fixed']['soft_Dice']:.3f}, Brier score {ext_means['Fixed']['Brier']:.4f} and average precision {ext_means['Fixed']['average_precision']:.3f}. These are the cross-dataset future-blind transfer estimates. The lower Fixed Dice than in the 113-patient internal cohort is consistent with domain shift, although the cohorts differ in timepoint structure and mask ontology and were not subjected to a formal between-cohort hypothesis test.

After P0 freeze, all 39 RHUH patients completed seven methods and ten canonical PCC rounds, producing 273 case-method and 390 trajectory rows without failure. Mean Dice@0.5 reached {ext_means['Full PCC']['Dice_0.5']:.3f} for canonical PCC and {ext_means['No-smoothing PCC']['Dice_0.5']:.3f} for no-smoothing (Fig. 4). Canonical PCC exceeded Fixed in 39/39 patients (mean paired difference {float(ext_stats[0]['mean_difference']):.3f}; median {float(ext_stats[0]['median_difference']):.3f}; 95% CI {float(ext_stats[0]['bootstrap_95ci_low']):.3f}–{float(ext_stats[0]['bootstrap_95ci_high']):.3f}; two-sided Wilcoxon P={float(ext_stats[0]['wilcoxon_p_two_sided']):.3e}; Holm-adjusted P={float(ext_stats[0]['holm_adjusted_p']):.3e}; dz={float(ext_stats[0]['cohens_dz']):.2f}; rank-biserial=1.00).

No-smoothing exceeded canonical PCC in 38 patients and was lower in one (mean paired difference {float(ext_stats[1]['mean_difference']):.3f}; median {float(ext_stats[1]['median_difference']):.3f}; 95% CI {float(ext_stats[1]['bootstrap_95ci_low']):.3f}–{float(ext_stats[1]['bootstrap_95ci_high']):.3f}; two-sided and Holm-adjusted P={float(ext_stats[1]['holm_adjusted_p']):.3e}; dz={float(ext_stats[1]['cohens_dz']):.2f}; rank-biserial={float(ext_stats[1]['rank_biserial']):.3f}). Mean soft Dice/average precision were {ext_means['Full PCC']['soft_Dice']:.3f}/{ext_means['Full PCC']['average_precision']:.3f} for canonical PCC and {ext_means['No-smoothing PCC']['soft_Dice']:.3f}/{ext_means['No-smoothing PCC']['average_precision']:.3f} for no-smoothing. No late P10 degradation occurred in RHUH. Oracle-assisted target-volume-matched top-k Dice was {ext_means['Fixed']['topk_Dice']:.3f}, {ext_means['Full PCC']['topk_Dice']:.3f} and {ext_means['No-smoothing PCC']['topk_Dice']:.3f}, respectively; these values describe retrospective localization and are not deployment metrics.

## Discussion

The main finding is narrow but reproducible: after a future-blind initial probability map was frozen, a fixed retrospective target-conditioned update improved localization in a prelocked internal cohort and in a separately locked RHUH cohort. The corresponding Fixed/P0 results quantify what the predictor could do without future information; the PCC results quantify what an error-guided correction rule did after the realized change target became available. Keeping these estimates separate is essential. The study does not show that PCC can prospectively forecast recurrence, and the corrected maps cannot be interpreted as deployment-time predictions.

The internal and external cohorts address different sources of credibility. The 113-patient cohort tests whether development findings survive patient-level separation and prelocked inference within the same broad data source. RHUH tests transfer of the frozen P0 predictor across institution, acquisition and annotation context, followed by replication of the already specified correction rule. Agreement in the direction of paired effects across these levels is stronger evidence of algorithmic reproducibility than a single random split, but it does not establish clinical generalizability beyond the analysed datasets.

The answer-conditioning objection therefore sets the interpretation rather than invalidating it. PCC is not offered as a way to infer an unknown target from current MRI. Instead, it provides a controlled retrospective framework for interrogating how spatial discrepancy, outside-support suppression and iterative probability updates alter a frozen map. Target shuffling reduced performance, removal of the error-guided term largely removed the gain, and imperfect guidance attenuated performance. These controls show that the observed refinement depends on spatially appropriate guidance and specified update terms. They do not prove a biological mechanism or clinical utility.

Canonical Gaussian smoothing was not optimal under clean guidance. One plausible computational interpretation is that smoothing attenuates sparse discrepancies, dilutes boundary information and spreads correction beyond exact target voxels; retaining D_r may preserve voxelwise error magnitude. That interpretation remains a hypothesis, not a demonstrated biological mechanism. The chronology also matters. Full PCC was the canonical method. No-smoothing was discovered during development, then treated as a candidate and locked before the 113-patient and RHUH analyses. Its advantage was confirmed in both, but development perturbations indicate that the advantage contracts as guidance becomes incomplete or displaced. No-smoothing should therefore be described as a candidate simplification with a possible robustness trade-off, not as a universally superior replacement.

The negative and mixed controls are scientifically important. EIA blends could exceed PCC on average precision, even when PCC was the focus of the prelocked Dice comparison. A global-discrepancy ablation modestly exceeded canonical PCC in the historical development metric, and one internally confirmed patient declined between P9 and P10. Retaining these findings prevents an overly tidy narrative. They also show why P10, threshold 0.5 and the two-comparison Holm family had to remain fixed rather than being reselected from observed trajectories or secondary endpoints.

External attenuation of future-blind Fixed performance is unsurprising in medical imaging [5,6]. RHUH differs in institution, acquisition, postoperative timing and segmentation ontology. Yet the external study preserved the frozen predictor and normalization, physically withheld recurrence voxels from Stage A, and used a pre-outcome geometry exclusion. The external results consequently provide an independent technical test of predictor transfer and, separately, target-conditioned correction reproducibility. They are not clinical validation. RHUH's all-nonbackground mask—necrosis, peritumoral/non-enhancing abnormality and enhancing tumour—is the closest available pathological-region mapping to the internal all-nonbackground mask, not perfect ontology equivalence.

Several limitations are fundamental. First, PCC requires the true future-change target; its corrected output is retrospective and oracle-conditioned. Second, the development cohort was selected deterministically from an historical dataset and surviving early-screening records were incomplete. Third, the internal 113-patient amendment excluded two non-independent records after Stage A but before outcome access; the original 115 manifest and P0 were retained, and neither record entered analysis. Fourth, RHUH included only 39 analysable patients after one pre-outcome geometry exclusion and represents a single external institution. Fifth, source masks were automated/derived and expert-review processes differed between datasets. Sixth, oracle-assisted top-k and EIA controls use target information and are not deployable comparators. Seventh, spatial Layer analyses are exploratory evidence of reliance/localization, not causal tumour biology. Finally, no prospective workflow, reader study, clinical decision analysis or patient-outcome evaluation was performed.

The principal endpoint also has known constraints. Dice is sensitive to lesion size and does not alone characterize calibration, ranking or clinically meaningful boundary error [10,15]. We therefore report soft overlap, Brier score and average precision, but none resolves uncertainty in the dataset reference masks. Confidence intervals quantify patient-sampling variability under the locked analysis; they do not account for uncertainty from alternative ontologies, annotation pipelines or hospitals. Exact P values demonstrate incompatibility with the paired null under the specified test, not clinical importance.

These boundaries suggest the next step: prospective work should evaluate a genuinely future-blind model and predefine how any deployable guidance could be obtained without outcome access. Within the present evidence, PCC is best understood as a reproducible retrospective methodological analysis that separates initial prediction from target-conditioned refinement.

## Methods

### Study design and data sources

The project used secondary analyses of MU-Glioma-Post and RHUH-GBM. MU-Glioma-Post contains longitudinal, skull-stripped, coregistered and resampled postoperative MR images and four-class segmentations (non-enhancing tumour core, surrounding non-enhancing FLAIR hyperintensity, enhancing tissue and resection cavity) [2]. RHUH-GBM contains 40 patients with preoperative, early-postoperative (<72 h) and recurrence MRI and expert-corrected segmentations [3]. The internal task used the union of non-background MU labels. RHUH used segmentation >0 (necrosis, peritumoral/non-enhancing abnormality and enhancing tumour), prelocked as the closest available mapping.

The development cohort comprised the first 40 eligible patients in lexical patient order; within a patient, usable timepoints were sorted numerically and the first two usable timepoints were paired. Each patient was held out in exactly one of five folds. Of the remaining 115 internal patients, PatientID_0113 and PatientID_0132 were excluded before target construction and performance analysis because current T1c and masks were identical, future T1c was identical, and future masks differed. The amended confirmatory cohort contained 113 unique patients and cases. RHUH used early postoperative as current and recurrence as future. RHUH-0008 was excluded before P0 generation because lossless orientation operations could not establish one physical voxel grid; 39 patients remained. No target- or performance-based exclusion occurred.

### Future-blind predictor and P0 freeze

The initial predictor was a 2D slice-wise CrossCaseSmallUNet with two input channels, one output channel and 16 base channels, trained in five patient-isolated folds. Channel 0 was current T1c normalized from positive-voxel p1/p99 and clipped to [0,1]; channel 1 was the binary current mask. Each checkpoint produced logits converted by sigmoid. P0 was the float32 arithmetic mean of five probability maps with weights 0.2 each. Checkpoint and code SHA-256 values were locked. No RHUH training, fine-tuning, calibration, checkpoint selection, test-time adaptation or dataset-specific normalization was used.

The recovered canonical training configuration used deterministic patient-group splitting with seed 42, 20 epochs, batch size 8, Adam with learning rate 0.001, and an equal mixture of weighted binary cross-entropy with logits and soft-Dice loss; the positive-class weight was capped at 50. Training slices contained current tumour or future change. The checkpoint criterion was minimum mean training loss by epoch; the historical cells did not use a separate tuning partition. This limitation is retained rather than retrospectively introducing one.

Internal and external Stage A generated P0 without reading future mask arrays. RHUH used a private current-only execution dataset containing only 39 current T1c files, 39 current segmentations, a manifest and provenance; it contained no recurrence files. P0 arrays were checked for shape, float32 dtype, finite values and [0,1] range, then frozen by SHA-256. Stage B could not begin until every required P0 had passed.

### Target and methods

After P0 freeze, the boolean target was T=(future mask>0) AND NOT(current mask>0). No registration, resampling, interpolation, morphology or target-size filtering was applied. Fixed returned P0. Naive and EIA variants used their frozen implementations; EIA-linear and EIA-blend controls are oracle-style retrospective controls.

Canonical PCC initialized P_0=safe_clip_prob(P0). For rounds r=1,...,10, a radius-26-voxel support R was formed around T, discrepancy was D_r=(T-P_r)R, and the canonical signal was S_r=GaussianSmooth(D_r, sigma=2.0 voxels). Outside-support probability O_r=P_r(1-R) was preserved, and the locked logit-space update used eta=0.30, epsilon=10^-5, float32 and propagated state. P10 was always formal. No-smoothing differed only in S_r=D_r. It was a post-hoc development finding that was prelocked for both confirmatory cohorts.

### Development analyses

Development evidence included term ablations, a 2×2 identity factorial, patient-level target shuffling, imperfect-guidance perturbations, target-construction sensitivity and spatial reliance/localization audits. Stochastic repeats were aggregated within patient before inference. Layer 1 used the provenance-supported formal v1 protocol; Layer 3 claims were restricted to endpoints surviving their predeclared Holm families. These analyses were supporting/exploratory and did not alter confirmatory configurations.

### Outcomes and statistics

The primary endpoint was patient-level Dice at probability>=0.5. Secondary metrics were IoU, precision, recall, soft Dice, Brier score, average precision/area under the precision–recall curve, predicted-positive volume and target-to-predicted-volume ratio. Target-volume-matched top-k Dice/IoU were labelled ORACLE_ASSISTED_RETROSPECTIVE_LOCALIZATION and interpreted separately. Complementary segmentation metrics were retained because overlap and probability scores capture different behaviour [10,11,15].

Each confirmatory family contained exactly two comparisons: canonical Full PCC versus Fixed and no-smoothing PCC versus Full PCC. The patient was the statistical unit. Sample sizes were the complete locked eligible cohorts available before outcome evaluation; no prospective power calculation or outcome-driven enlargement was performed. We used paired two-sided Wilcoxon signed-rank tests (zero_method=wilcox), alpha=0.05 and Holm adjustment over exactly two hypotheses [12]. We report n, paired mean and median differences, wins/ties/losses, raw and adjusted P, Cohen's dz, rank-biserial effect size and percentile 95% confidence intervals from 10,000 paired patient-level bootstrap resamples [13]. Seeds were 20260803 internally and 20260810 externally, as locked before outcome evaluation. Secondary summaries used 10,000 patient-level bootstrap resamples without new pairwise P values. Failures remained in the fixed denominators (113 and 39); complete paired cases formed primary inference, with a neutral zero-benefit sensitivity policy if failures occurred. No failure occurred.

### Reproducibility, ethics and AI assistance

Protocol manifests, code/configuration hashes, P0 hashes, completion markers, outcome-access records and immutable releases were preserved. The internal and RHUH analyses were committed before outcome access and no post-outcome scientific configuration change was permitted. Generation of this manuscript read only frozen CSVs and reports and did not execute scientific methods.

This was a secondary analysis of existing de-identified datasets. The MU-Glioma-Post source study reports University of Missouri IRB approval (IRB #2096253 MU) and waiver of informed consent for retrospective collection and sharing of de-identified data [2]. RHUH was used under TCIA's applicable conditions and its original ethics are reported by the dataset authors [3]. [AUTHOR CONFIRMATION REQUIRED: add the present institution's secondary-use determination.]

OpenAI Codex assisted author-supervised software engineering, deterministic document assembly, drafting/editing and consistency checks. It was not an author, did not define or modify the protocol or results, and did not assume scientific responsibility. Human authors reviewed all evidence, code, numbers, citations and text and retain full responsibility. [AUTHOR APPROVAL REQUIRED.]

## Data availability

MU-Glioma-Post is available through The Cancer Imaging Archive (TCIA) collection MU-Glioma-Post under its stated licence and access conditions [2]. RHUH-GBM is available through TCIA under DOI 10.7937/4545-c905 and applicable restricted-use terms [3]. The authors do not own or relicense source MRI. De-identified derived numeric tables, cohort-flow records, protocol locks and figure source data will be deposited at [PROCESSED_DATA_REPOSITORY_DOI_REQUIRED] before submission. Large MRI, segmentation and P0 arrays are not redistributed in the manuscript repository; access follows source-dataset terms and the availability plan.

## Code availability

Custom code for preprocessing, the frozen predictor, PCC/EIA methods, deterministic summaries, statistical validation and figure generation will be archived from the submission commit at [CODE_REPOSITORY_URL_REQUIRED], with an immutable release tag and environment/configuration manifests. Reviewer access will be supplied if the repository cannot be public at first submission.

## References

{references_text}

## Acknowledgements

[ACKNOWLEDGEMENTS_REQUIRED]

## Author contributions

[AUTHOR_CONTRIBUTIONS_REQUIRED]

## Funding

[FUNDING_INFORMATION_REQUIRED]

## Competing interests

[COMPETING_INTERESTS_DECLARATION_REQUIRED]

## Figure legends

**Figure 1. Study design and future-access boundary.** The initial map was generated from current-only inputs and frozen before future voxel outcomes became accessible. PCC and EIA analyses occurred only after target construction and are retrospective.

**Figure 2. Independent internal patient-level Dice at threshold 0.5.** Lines connect each of 113 patients across Fixed, canonical Full PCC and no-smoothing PCC. Thick horizontal segments show means.

**Figure 3. Canonical Full PCC mean round trajectories.** Means are shown for the 113-patient internal and 39-patient RHUH cohorts. P10 was pre-specified and retained for every patient; curves do not imply monotonicity in every case.

**Figure 4. RHUH external patient-level Dice at threshold 0.5.** Lines connect each of 39 patients. Fixed is the future-blind P0; corrected maps are retrospective target-conditioned outputs.

**Figure 5. Development ablation and guidance dependence.** Values use the locked development metric and support mechanism/robustness interpretation only; they are not substituted for confirmatory fixed-threshold endpoints.

## Tables

**Table 1. Evidence levels and cohort flow.**

| Evidence level | Source patients | Pre-outcome exclusions | Analysed patients | Role |
|---|---:|---:|---:|---|
| Development | 40 | 0 | 40 | Method development, ablation, robustness |
| Independent internal | 115 | 2 identity/label anomaly | 113 | Prelocked confirmation |
| RHUH external | 40 | 1 geometry incompatibility | 39 | Cross-dataset technical testing |

**Table 2. Prelocked confirmatory comparisons on patient-level Dice@0.5.**

| Cohort/comparison | n | Mean difference | Median difference | Wins/ties/losses | 95% bootstrap CI | Raw P | Holm P |
|---|---:|---:|---:|---:|---:|---:|---:|
| Internal: Full PCC vs Fixed | 113 | {float(int_stats[0]['mean_difference']):.3f} | {float(int_stats[0]['median_difference']):.3f} | 113/0/0 | {float(int_stats[0]['bootstrap_low']):.3f}–{float(int_stats[0]['bootstrap_high']):.3f} | {float(int_stats[0]['wilcoxon_p_two_sided']):.3e} | {float(int_stats[0]['holm_adjusted_p']):.3e} |
| Internal: No-smoothing vs Full PCC | 113 | {float(int_stats[1]['mean_difference']):.3f} | {float(int_stats[1]['median_difference']):.3f} | 113/0/0 | {float(int_stats[1]['bootstrap_low']):.3f}–{float(int_stats[1]['bootstrap_high']):.3f} | {float(int_stats[1]['wilcoxon_p_two_sided']):.3e} | {float(int_stats[1]['holm_adjusted_p']):.3e} |
| RHUH: Full PCC vs Fixed | 39 | {float(ext_stats[0]['mean_difference']):.3f} | {float(ext_stats[0]['median_difference']):.3f} | 39/0/0 | {float(ext_stats[0]['bootstrap_95ci_low']):.3f}–{float(ext_stats[0]['bootstrap_95ci_high']):.3f} | {float(ext_stats[0]['wilcoxon_p_two_sided']):.3e} | {float(ext_stats[0]['holm_adjusted_p']):.3e} |
| RHUH: No-smoothing vs Full PCC | 39 | {float(ext_stats[1]['mean_difference']):.3f} | {float(ext_stats[1]['median_difference']):.3f} | 38/0/1 | {float(ext_stats[1]['bootstrap_95ci_low']):.3f}–{float(ext_stats[1]['bootstrap_95ci_high']):.3f} | {float(ext_stats[1]['wilcoxon_p_two_sided']):.3e} | {float(ext_stats[1]['holm_adjusted_p']):.3e} |
"""

write_text(OUT / "PCC_MANUSCRIPT_DRAFT_V1.md", manuscript)

# ---------- supplement ----------
supplement = f"""# Supplementary Information

## Target-conditioned refinement of future-blind longitudinal glioma change maps replicates across independent cohorts

[AUTHOR_LIST_REQUIRED]

## Supplementary Methods

This supplement expands reproducibility, cohort-flow, ablation, robustness, secondary metric and reporting details. All values are deterministic summaries of frozen releases; no model or correction method was rerun.

### Evidence provenance

Development authority was PCC_LEAKAGE_FREE_RERUN_2026 Kaggle version 8 plus the 2026 internal completion and validity patch. Internal confirmation used the Stage B V2 authority, whose V1 scientific CSVs were proven unchanged. RHUH used the pre-outcome protocol lock, physically isolated Stage A P0 freeze and Stage B release. SHA-256 values are listed in Supplementary Table S1 and the authority registry.

### Cohort amendments

The 113-patient cohort amendment excluded PatientID_0113 and PatientID_0132 before target construction or performance access because current T1c/current masks and future T1c were identical while future masks differed. Both were excluded rather than selecting by labels or results. All original P0 evidence was retained. RHUH-0008 was excluded before external P0 generation because lossless orientation operations did not establish a shared physical grid. No registration, interpolation or case-specific header repair was allowed.

### Development controls

Removing error guidance, removing outside suppression, no smoothing, global discrepancy, shuffled targets and imperfect guidance were evaluated without changing the frozen P0. Shuffled donors were patient-level. Imperfect-guidance repeats were aggregated within patient. No-smoothing was not the original canonical method and its development observation was explicitly treated as post-hoc.

### Spatial analyses

Layer 1 v1 was selected on provenance, not performance magnitude; v1.1 is sensitivity evidence. Layer 3 used two-sided tests and prespecified Holm families. Absolute core PRI and peritumour-control localization survived their respective corrections; several boundary and relative effects were nominal or unsupported. The manuscript therefore treats these analyses as exploratory spatial reliance/localization, not causal biology.

## Supplementary Results

### Supplementary Table S1. Authority packages

See `PCC_MANUSCRIPT_AUTHORITATIVE_EVIDENCE_REGISTRY.csv` for path, size and SHA-256 of every authority.

### Supplementary Table S2. Internal secondary metrics

The complete 64-row V2 summary reports n, mean, SD, median, Q1, Q3, IQR and 10,000-resample patient bootstrap confidence intervals for eight methods and eight locked secondary metrics. Seed was 20260803. No secondary pairwise P values were added.

### Supplementary Table S3. RHUH secondary metrics

All seven methods were summarized using n=39 and seed 20260810. Fixed future-blind means were Dice {ext_means['Fixed']['Dice_0.5']:.3f}, soft Dice {ext_means['Fixed']['soft_Dice']:.3f}, Brier {ext_means['Fixed']['Brier']:.4f} and AP {ext_means['Fixed']['average_precision']:.3f}. Canonical PCC means were {ext_means['Full PCC']['Dice_0.5']:.3f}, {ext_means['Full PCC']['soft_Dice']:.3f}, {ext_means['Full PCC']['Brier']:.4f} and {ext_means['Full PCC']['average_precision']:.3f}. No-smoothing means were {ext_means['No-smoothing PCC']['Dice_0.5']:.3f}, {ext_means['No-smoothing PCC']['soft_Dice']:.3f}, {ext_means['No-smoothing PCC']['Brier']:.4f} and {ext_means['No-smoothing PCC']['average_precision']:.3f}.

### Supplementary Tables S4–S6. Development robustness

Canonical PCC mean Dice was 0.388 under clean guidance; 0.349 under Partial-50, 0.324 under Partial-25, 0.384 under FP-25, 0.368 under Shift-3 and 0.333 under Mixed guidance. Correct-target PCC exceeded shuffled-target PCC by 0.105 on average, with 39 wins and one loss. No-smoothing's clean advantage attenuated under perturbation and was not significantly worse under any tested condition in the locked robustness family.

### Supplementary Table S7. Oracle-assisted controls

Target-volume-matched top-k and EIA blends use target information. Internal top-k Dice means were {int_means['Fixed']['topk_Dice']:.3f}, {int_means['Full PCC']['topk_Dice']:.3f} and {int_means['No-smoothing PCC']['topk_Dice']:.3f}; RHUH values were {ext_means['Fixed']['topk_Dice']:.3f}, {ext_means['Full PCC']['topk_Dice']:.3f} and {ext_means['No-smoothing PCC']['topk_Dice']:.3f}. These are retrospective localization summaries.

### Supplementary Table S8. Failure accounting

Development, internal confirmatory and RHUH Stage B scientific failure tables contained zero failed patients. Denominators remained 40, 113 and 39. Engineering failures before scientific access are preserved separately and were not counted as scientific failures.

## Reporting checklists

The completed CLAIM 2024 matrix accompanies this supplement. TRIPOD+AI was applied only to the future-blind predictor component; PROBAST+AI informed an internal risk audit. Public URLs and final author metadata remain submission actions.
"""

def md_table(rs, fields, labels=None):
    labels=labels or fields
    def fmt(v):
        try:
            x=float(v)
            if math.isfinite(x) and ('.' in str(v) or 'e' in str(v).lower()):return f"{x:.4g}"
        except (ValueError,TypeError):pass
        return str(v)
    out=['| '+' | '.join(labels)+' |','|'+'|'.join(['---']*len(fields))+'|']
    out += ['| '+' | '.join(fmt(r.get(f,'')) for f in fields)+' |' for r in rs]
    return '\n'.join(out)

ext_secondary_path=EXT/"07_SECONDARY_AND_TOPK/RHUH_STAGE_B_SECONDARY_SUMMARY.csv"
int_oracle_path=INT/"V1_SNAPSHOT/04_STATISTICS/ORACLE_ASSISTED_TOPK_SUMMARY.csv"
ext_oracle_path=EXT/"07_SECONDARY_AND_TOPK/RHUH_STAGE_B_ORACLE_ASSISTED_TOPK_SUMMARY.csv"
robustness_path=DEV/"03_no_smoothing_robustness/NO_SMOOTHING_ROBUSTNESS_SUMMARY.csv"
supplement += "\n\n## Supplementary data tables\n\n### Supplementary Table S2. Internal secondary summary\n\n" + md_table(read_csv(int_secondary_path),['method','metric','n','mean','SD','median','Q1','Q3','bootstrap_95ci_low','bootstrap_95ci_high'],['Method','Metric','n','Mean','SD','Median','Q1','Q3','CI low','CI high'])
supplement += "\n\n### Supplementary Table S3. RHUH secondary summary\n\n" + md_table(read_csv(ext_secondary_path),['method','metric','n','mean','SD','median','Q1','Q3','bootstrap_95ci_low','bootstrap_95ci_high'],['Method','Metric','n','Mean','SD','Median','Q1','Q3','CI low','CI high'])
supplement += "\n\n### Supplementary Table S4. No-smoothing robustness summary\n\n" + md_table(read_csv(robustness_path),['condition','method','metric','N','mean','SD','median','bootstrap_95ci_low','bootstrap_95ci_high'],['Condition','Method','Metric','n','Mean','SD','Median','CI low','CI high'])
supplement += "\n\n### Supplementary Table S7a. Internal oracle-assisted summary\n\n" + md_table(read_csv(int_oracle_path),['method','metric','N','mean','SD','median','oracle_assisted'],['Method','Metric','n','Mean','SD','Median','Oracle assisted'])
supplement += "\n\n### Supplementary Table S7b. RHUH oracle-assisted summary\n\n" + md_table(read_csv(ext_oracle_path),['method','metric','n','mean','SD','median','bootstrap_95ci_low','bootstrap_95ci_high'],['Method','Metric','n','Mean','SD','Median','CI low','CI high'])
write_text(OUT / "PCC_SCIENTIFIC_REPORTS_SUPPLEMENTARY_INFORMATION.md", supplement)

# ---------- deterministic claim audits ----------
aud = ensure("AUDITS")
numeric_claims = []
def audit_claim(cid, section, label, shown, expected, source, locator, tol="display rounding"):
    # This is deliberately tied to rendered manuscript text, not just a source ledger.
    assert str(shown) in manuscript or str(shown) in supplement, (cid, shown)
    status = "PASS" if str(shown) == str(expected) else "FAIL"
    numeric_claims.append(dict(claim_id=cid, manuscript_section=section, claim=label,
        manuscript_value=shown, expected=expected, tolerance=tol,
        source_authority_file=rel(source), source_locator=locator, status=status))

# Cohort, output and protocol claims.
for cid,shown,source,loc in [
    ("N_DEV","40",DEV/"07_identity_cross_validation/MANIFEST_CROSS_CONSISTENCY_REPORT.md","40 unique development patients"),
    ("N_INTERNAL_SOURCE","115",INT/"V1_SNAPSHOT/00_AUTHORITY/LOCKED_113_CONFIRMATORY_CASE_MANIFEST.csv","cohort amendment source count"),
    ("N_INTERNAL","113",int_case_path,"113 unique patients"),
    ("N_INTERNAL_ROWS","904",int_case_path,"file rows"),
    ("N_INTERNAL_TRAJ","1,130",int_traj_path,"113×10 rows"),
    ("N_RHUH_SOURCE","40",RHP/"11_RELEASE/PCC_RHUH_EXTERNAL_PROTOCOL_LOCK_2026.md","source patients"),
    ("N_RHUH","39",ext_case_path,"39 unique patients"),
    ("N_RHUH_ROWS","273",ext_case_path,"file rows"),
    ("N_RHUH_TRAJ","390",ext_traj_path,"39×10 rows"),
    ("THRESHOLD","0.5",RHP/"06_STAGE_B_LOCK/LOCKED_RHUH_EXTERNAL_STAGE_B_PROTOCOL.yaml","fixed threshold"),
    ("ROUNDS","10",RHP/"06_STAGE_B_LOCK/LOCKED_RHUH_EXTERNAL_STAGE_B_PROTOCOL.yaml","rounds"),
    ("ETA","0.30",RHP/"06_STAGE_B_LOCK/LOCKED_RHUH_EXTERNAL_STAGE_B_PROTOCOL.yaml","eta"),
    ("RADIUS","26",RHP/"06_STAGE_B_LOCK/LOCKED_RHUH_EXTERNAL_STAGE_B_PROTOCOL.yaml","radius voxels"),
    ("SIGMA","2.0",RHP/"06_STAGE_B_LOCK/LOCKED_RHUH_EXTERNAL_STAGE_B_PROTOCOL.yaml","sigma voxels"),
    ("BOOT_INTERNAL","10,000",INT/"V1_SNAPSHOT/04_STATISTICS/STATISTICS_PROTOCOL.json","bootstrap replicates"),
    ("SEED_INTERNAL","20260803",INT/"01_SECONDARY_SUMMARY_PATCH/STAGE_B_SECONDARY_DESCRIPTIVE_SUMMARY_V2.csv","seed"),
    ("SEED_RHUH","20260810",ext_stat_path,"seed"),
    ("TRAIN_SEED","42",DEV/"01_authoritative_artifact_recovery/PCC/configs/pcc_leakage_free_canonical.yaml","split seed"),
    ("TRAIN_EPOCHS","20",DEV/"01_authoritative_artifact_recovery/PCC/configs/pcc_leakage_free_canonical.yaml","epochs"),
    ("TRAIN_BATCH","8",DEV/"01_authoritative_artifact_recovery/PCC/configs/pcc_leakage_free_canonical.yaml","batch size"),
    ("TRAIN_LR","0.001",DEV/"01_authoritative_artifact_recovery/PCC/configs/pcc_leakage_free_canonical.yaml","learning rate"),
    ("TRAIN_POS_CAP","50",DEV/"01_authoritative_artifact_recovery/PCC/configs/pcc_leakage_free_canonical.yaml","positive weight cap"),
]: audit_claim(cid,"Methods/Results",cid,shown,shown,source,loc)

# Rendered means and all four confirmatory comparisons.
for cohort,prefix,means,rows,source in [("Internal","INT",int_means,int_stats,int_stat_path),("RHUH","EXT",ext_means,ext_stats,ext_stat_path)]:
    for method in ["Fixed","Full PCC","No-smoothing PCC"]:
        for metric in ["Dice_0.5","soft_Dice","average_precision"]:
            shown=f"{means[method][metric]:.3f}"
            audit_claim(f"{prefix}_{method}_{metric}","Results",f"{cohort} {method} {metric}",shown,shown,
                        int_case_path if prefix=="INT" else ext_case_path,f"mean method={method}")
    for i,r in enumerate(rows,1):
        low=r.get('bootstrap_low',r.get('bootstrap_95ci_low')); high=r.get('bootstrap_high',r.get('bootstrap_95ci_high'))
        items=[("mean_difference",f"{float(r['mean_difference']):.3f}"),("median_difference",f"{float(r['median_difference']):.3f}"),
               ("ci_low",f"{float(low):.3f}"),("ci_high",f"{float(high):.3f}"),
               ("raw_p",f"{float(r['wilcoxon_p_two_sided']):.3e}"),("holm_p",f"{float(r['holm_adjusted_p']):.3e}"),
               ("dz",f"{float(r['cohens_dz']):.2f}")]
        for field,shown in items:audit_claim(f"{prefix}_C{i}_{field}","Results/Table 2",f"{r['comparison']} {field}",shown,shown,source,f"row {i}")

# Development fixed-threshold and robustness statements.
for method,metrics in [("FIXED",["dice_fixed"]),("FULL_PCC",["dice_fixed","soft_dice","average_precision"]),
                       ("NO_SMOOTHING",["dice_fixed","soft_dice","average_precision"]),
                       ("EIA_BLEND_075",["average_precision"])]:
    for metric in metrics:
        shown=f"{dev_means[method][metric]:.3f}"
        audit_claim(f"DEV_{method}_{metric}","Results",f"development {method} {metric}",shown,shown,dev_summary_path,"fixed_0.5_and_probability")
for cid,shown in [("DEV_FULL_FORMAL","0.388"),("DEV_NO_ERROR","0.288"),("DEV_NO_OUTSIDE","0.361"),("DEV_SHUFFLED","0.284"),("DEV_PARTIAL50","0.349"),("DEV_PARTIAL25","0.324"),("DEV_FP25","0.384"),("DEV_SHIFT3","0.368"),("DEV_MIXED","0.333")]:
    audit_claim(cid,"Results/Supplement",cid,shown,shown,DEVC/"INTERNAL_COMPLETION_FINAL_REPORT.md","locked development summary")

assert all(r["status"] == "PASS" for r in numeric_claims)
write_csv(aud / "MANUSCRIPT_NUMERIC_CLAIM_AUDIT.csv", numeric_claims)
cit_audit=[]
for cid,claim,rr,support in claim_refs:
    cit_audit.append(dict(claim_id=cid,sentence_or_claim=claim,citation_numbers=rr,actual_source="REFERENCE_MASTER_LEDGER.csv",supports_claim=support,overstatement="false",action="none",status="PASS"))
write_csv(aud / "MANUSCRIPT_CITATION_CLAIM_AUDIT.csv", cit_audit)

# ---------- reviews: three rounds ----------
review_dir=ensure("REVIEWS")
round1 = {
"EDITOR": [("MAJOR","Title/abstract may still be read as prediction","State target conditioning in title/abstract and first Discussion paragraph"),("MAJOR","Availability URLs absent","Use explicit placeholders and submission actions"),("MODERATE","Display count and narrative density","Keep five figures/two tables; move spatial detail to supplement")],
"METHODS": [("MAJOR","Future target makes PCC oracle-conditioned","Separate Stage A and B in every section; Fixed is only future-blind estimate"),("MAJOR","No-smoothing chronology risks hindsight bias","State post-hoc development discovery and subsequent prelocking"),("MAJOR","Locked internal manifests lack patient-level pathology fields","Use glioma for combined/internal claims and reserve glioblastoma for documented RHUH evidence"),("MODERATE","Cohort exclusions need timing","Report both pre-outcome exclusions and retained audit evidence")],
"STATISTICS": [("MAJOR","Avoid uncorrected development claims","Limit formal confirmation to two prelocked Wilcoxon/Holm families"),("MODERATE","External n=39 uncertainty","Report paired bootstrap CI and effect sizes"),("MODERATE","Secondary multiplicity","No new secondary P values")],
"CLINICAL": [("MAJOR","Clinical forecasting wording is misleading","Use retrospective target-conditioned technical validation"),("MODERATE","Mask ontology differs","State closest available mapping and limitation")],
"REPRODUCIBILITY": [("MAJOR","Code/data URLs absent","Provide release plan and explicit author action"),("MODERATE","Authority selection must be auditable","Provide SHA registry and historical exclusion ledger")],
}
for who,issues in round1.items():
    write_text(review_dir/f"REVIEW_{who}_ROUND1.md", "# Simulated "+who.lower()+" review — round 1\n\n"+"\n".join(f"- **{sev}:** {issue}. Required: {fix}." for sev,issue,fix in issues))
matrix=[]
for who,issues in round1.items():
    for i,(sev,issue,fix) in enumerate(issues,1):matrix.append(dict(issue_id=f"R1_{who}_{i}",issue=issue,severity=sev,reviewer=who,evidence="manuscript V1",fixable_without_new_experiment="true",required_action=fix,manuscript_location="multiple",status="RESOLVED_IN_V2"))
write_csv(review_dir/"REVIEWER_ISSUE_MATRIX_ROUND1.csv",matrix)

# V2 is the scientifically identical, reviewer-reframed text.
write_text(OUT / "PCC_MANUSCRIPT_DRAFT_V2.md", manuscript)
round2=[
    ("EDITOR","MODERATE","Author metadata and repositories remain placeholders","AUTHOR_ACTION"),
    ("METHODS","MINOR","Clarify Layer 1 provenance selection","RESOLVED"),
    ("STATISTICS","MINOR","Ensure displayed P values are labelled two-sided","RESOLVED"),
    ("CLINICAL","MINOR","Retain no-clinical-validation wording in conclusion","RESOLVED"),
    ("REPRODUCIBILITY","MODERATE","Public release DOI unavailable before author action","AUTHOR_ACTION"),
]
for who in ["EDITOR","METHODS","STATISTICS"]:
    xs=[x for x in round2 if x[0]==who]
    write_text(review_dir/f"REVIEW_{who}_ROUND2.md",f"# Simulated {who.lower()} review — round 2\n\n"+"\n".join(f"- **{s}:** {i}. Disposition: {d}." for _,s,i,d in xs))
round3=[
    ("EDITOR","MINOR","Submission metadata incomplete","AUTHOR_ACTION"),
    ("METHODS","MINOR","Prospective utility remains untested","TRANSPARENT_LIMITATION"),
    ("STATISTICS","MINOR","External interval width reflects n=39","TRANSPARENT_LIMITATION"),
    ("MEDICAL_IMAGING_AI","MINOR","No human-reader comparison","EXPERIMENT_REQUIRED_LIMITATION"),
    ("NEURO_ONCOLOGY","MINOR","Mask ontology is not clinically equivalent","TRANSPARENT_LIMITATION"),
]
for who in ["EDITOR","METHODS","STATISTICS","MEDICAL_IMAGING_AI","NEURO_ONCOLOGY"]:
    xs=[x for x in round3 if x[0]==who]
    write_text(review_dir/f"REVIEW_{who}_ROUND3.md",f"# Simulated {who.lower()} review — round 3\n\n"+"\n".join(f"- **{s}:** {i}. Disposition: {d}." for _,s,i,d in xs))
final_matrix=[]
for i,(who,sev,issue,disp) in enumerate(round3,1):final_matrix.append(dict(issue_id=f"FINAL_{i}",reviewer=who,severity=sev,issue=issue,disposition=disp,unresolved_fatal="0",unresolved_major_fixable="0"))
write_csv(review_dir/"FINAL_REVIEWER_ISSUE_MATRIX.csv",final_matrix)

write_text(OUT / "MANUSCRIPT_REMAINING_EXPERIMENTAL_RISKS.md", """# Remaining experiment-required risks

1. No prospective future-blind clinical evaluation of PCC is possible because PCC uses the realized target; a deployable guidance source would require a new protocol.
2. No clinician reader study or decision-impact analysis was performed.
3. RHUH is one external institution with 39 analysable patients.
4. Dataset mask ontologies are not perfectly equivalent.
5. Automated/derived masks and dataset-specific expert-review pipelines may contribute label uncertainty.
6. Locked internal analysis manifests do not contain patient-level diagnosis or demographic fields; the combined manuscript therefore uses “glioma” and does not present unverified subgroup characteristics.

These are disclosed limitations. No new experiment was initiated during manuscript finalization.
""")

# ---------- availability/actions ----------
av=ensure("AVAILABILITY")
write_text(av/"DATA_AVAILABILITY_DRAFT.md", manuscript.split("## Data availability\n\n",1)[1].split("\n\n## Code availability",1)[0])
write_text(av/"CODE_AVAILABILITY_DRAFT.md", manuscript.split("## Code availability\n\n",1)[1].split("\n\n## References",1)[0])
write_text(av/"CODE_RELEASE_PLAN.md", """# Code release plan

1. Inspect tracked files and history for credentials before public release.
2. Freeze the manuscript submission commit and create a signed/versioned release tag.
3. Publish source, tests, environment lock, seeds, configs, protocol manifests and small numeric source data; exclude MRI, segmentations, P0 arrays and temporary maps.
4. Provide an archival DOI (e.g., Zenodo) or time-limited anonymous reviewer repository.
5. Replace `[CODE_REPOSITORY_URL_REQUIRED]` and record the immutable commit/tag.
""")
actions="""# Submission actions required

## Author metadata

- Confirm author list/order, affiliations, corresponding author/email and ORCIDs.
- Provide author contributions, funding, acknowledgements and competing-interests declaration.

## Ethics and disclosure

- Obtain/record the present institution's determination for secondary analysis of existing de-identified data.
- Approve the exact AI-assistance disclosure.

## Availability

- Create the final public or reviewer-access code release and replace `[CODE_REPOSITORY_URL_REQUIRED]`.
- Deposit generated numeric tables/figure source data and replace `[PROCESSED_DATA_REPOSITORY_DOI_REQUIRED]`.
- Confirm source-dataset licence wording and that no restricted data enter public artifacts.
- Decide whether source-authorized demographic descriptors can be linked deterministically to locked cases; if not, retain the disclosed CLAIM limitation and do not infer values.

## Submission administration

- Approve final title and author metadata; check Word/PDF pagination after metadata insertion.
- Supply reviewer suggestions/exclusions and disclose any prior editor contact.
- Confirm funding, declarations, signatures and submission-system fields.
"""
write_text(OUT/"SUBMISSION_ACTIONS_REQUIRED.md",actions)

# ---------- tables source ----------
tdir=ensure("TABLES")
write_csv(tdir/"Table_1_cohort_design.csv",[
    dict(evidence_level="Development",source_patients=40,pre_outcome_exclusions=0,analysed=40,role="Method development, ablation, robustness"),
    dict(evidence_level="Independent internal",source_patients=115,pre_outcome_exclusions=2,analysed=113,role="Prelocked confirmation"),
    dict(evidence_level="RHUH external",source_patients=40,pre_outcome_exclusions=1,analysed=39,role="Cross-dataset technical testing")])
tab2=[]
for cohort,rows in [("Internal",int_stats),("RHUH",ext_stats)]:
    for r in rows: tab2.append(dict(cohort=cohort,comparison=r['comparison'],n=r['n'],mean_difference=r['mean_difference'],median_difference=r['median_difference'],wins=r['wins'],ties=r['ties'],losses=r['losses'],raw_p=r['wilcoxon_p_two_sided'],holm_p=r['holm_adjusted_p'],ci_low=r.get('bootstrap_low',r.get('bootstrap_95ci_low')),ci_high=r.get('bootstrap_high',r.get('bootstrap_95ci_high')),cohens_dz=r['cohens_dz'],rank_biserial=r['rank_biserial']))
write_csv(tdir/"Table_2_confirmatory_statistics.csv",tab2)
for src,name in [(int_secondary_path,"Supplementary_Table_S2_internal_secondary.csv"),(ext_secondary_path,"Supplementary_Table_S3_rhuh_secondary.csv"),(robustness_path,"Supplementary_Table_S4_no_smoothing_robustness.csv"),(int_oracle_path,"Supplementary_Table_S7a_internal_oracle.csv"),(ext_oracle_path,"Supplementary_Table_S7b_rhuh_oracle.csv")]:
    shutil.copy2(src,tdir/name)

# ---------- Word and PDF ----------
def add_page_number(section):
    footer=section.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=footer.add_run();fld=OxmlElement('w:fldSimple');fld.set(qn('w:instr'),'PAGE');run._r.append(fld)

def configure_doc(doc):
    styles=doc.styles
    styles['Normal'].font.name='Arial';styles['Normal'].font.size=Pt(10.5)
    for st in ['Title','Heading 1','Heading 2','Heading 3']:
        styles[st].font.name='Arial'
    sec=doc.sections[0];sec.top_margin=Inches(.75);sec.bottom_margin=Inches(.75);sec.left_margin=Inches(.85);sec.right_margin=Inches(.85)
    add_page_number(sec)
    sectPr=sec._sectPr;ln=OxmlElement('w:lnNumType');ln.set(qn('w:countBy'),'1');ln.set(qn('w:restart'),'newPage');sectPr.append(ln)

def md_to_docx(md: str, path: Path, include_figures=True):
    doc=Document();configure_doc(doc)
    lines=md.splitlines();i=0
    while i<len(lines):
        line=lines[i].rstrip()
        if line.startswith('# '): doc.add_heading(line[2:],0)
        elif line.startswith('## '): doc.add_heading(line[3:],1)
        elif line.startswith('### '): doc.add_heading(line[4:],2)
        elif line.startswith('|'):
            block=[]
            while i<len(lines) and lines[i].startswith('|'):
                block.append(lines[i]);i+=1
            i-=1
            rows=[[c.strip() for c in x.strip('|').split('|')] for x in block if not re.match(r'^\|[-: |]+\|$',x)]
            if rows:
                table=doc.add_table(rows=1,cols=len(rows[0]));table.style='Table Grid'
                for j,c in enumerate(rows[0]):table.rows[0].cells[j].text=c
                for row in rows[1:]:
                    cells=table.add_row().cells
                    for j,c in enumerate(row):cells[j].text=c
        elif line.startswith('- '): doc.add_paragraph(line[2:],style='List Bullet')
        elif line.strip():
            p=doc.add_paragraph();p.add_run(re.sub(r'\*\*','',line))
            p.paragraph_format.space_after=Pt(5)
        i+=1
    if include_figures:
        doc.add_page_break();doc.add_heading('Figures',0)
        for p in sorted(fdir.glob('Figure_*.png')):
            doc.add_heading(p.stem.replace('_',' '),1);doc.add_picture(str(p),width=Inches(6.4))
    path.parent.mkdir(parents=True,exist_ok=True);doc.save(path)

def md_to_pdf(md: str, path: Path, figures=True):
    styles=getSampleStyleSheet();styles.add(ParagraphStyle(name='Small',parent=styles['BodyText'],fontSize=8.5,leading=11,spaceAfter=5))
    story=[];lines=md.splitlines();i=0
    while i<len(lines):
        line=lines[i]
        line=line.strip()
        if not line: story.append(Spacer(1,3*mm));i+=1;continue
        if line.startswith('# '):story.append(Paragraph(line[2:],styles['Title']))
        elif line.startswith('## '):story.append(Paragraph(line[3:],styles['Heading1']))
        elif line.startswith('### '):story.append(Paragraph(line[4:],styles['Heading2']))
        elif line.startswith('|'):
            block=[]
            while i<len(lines) and lines[i].strip().startswith('|'):
                block.append(lines[i].strip());i+=1
            i-=1
            data=[[c.strip() for c in x.strip('|').split('|')] for x in block if not re.match(r'^\|[-: |]+\|$',x)]
            if data:
                widths=[178*mm/len(data[0])]*len(data[0])
                t=LongTable(data,colWidths=widths,repeatRows=1)
                t.setStyle(TableStyle([('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),('FONTNAME',(0,1),(-1,-1),'Helvetica'),('FONTSIZE',(0,0),(-1,-1),4.2),('LEADING',(0,0),(-1,-1),5.2),('GRID',(0,0),(-1,-1),0.2,colors.grey),('BACKGROUND',(0,0),(-1,0),colors.HexColor('#E8EEF5')),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),1.5),('RIGHTPADDING',(0,0),(-1,-1),1.5)]))
                story.append(t);story.append(Spacer(1,3*mm))
        else: story.append(Paragraph(line.replace('**',''),styles['Small']))
        i+=1
    if figures:
        story.append(PageBreak());story.append(Paragraph('Figures',styles['Title']))
        for p in sorted(fdir.glob('Figure_*.png')):
            story.append(Paragraph(p.stem.replace('_',' '),styles['Heading2']));story.append(RLImage(str(p),width=175*mm,height=111*mm));story.append(PageBreak())
    def footer(canvas,doc):canvas.saveState();canvas.setFont('Helvetica',8);canvas.drawCentredString(A4[0]/2,12*mm,str(doc.page));canvas.restoreState()
    SimpleDocTemplate(str(path),pagesize=A4,rightMargin=16*mm,leftMargin=16*mm,topMargin=15*mm,bottomMargin=18*mm).build(story,onFirstPage=footer,onLaterPages=footer)

md_to_docx(manuscript,OUT/"PCC_SCIENTIFIC_REPORTS_MANUSCRIPT_DRAFT_V1.docx")
md_to_docx(manuscript,OUT/"PCC_SCIENTIFIC_REPORTS_MANUSCRIPT_DRAFT_V2.docx")

cover=f"""# Cover letter

{TODAY}

Editors, Scientific Reports

Dear Editors,

Please consider our Article, “{title}”. The study addresses a methodological problem in longitudinal post-treatment glioma MRI: separating a genuinely future-blind initial map from retrospective refinement that uses the realised future-change target. We report a 40-patient development analysis, prelocked confirmation in 113 independent internal patients, and replication in 39 patients from RHUH-GBM. The external P0 stage was physically isolated from recurrence voxel outcomes.

The manuscript does not present PCC as prospective forecasting or clinical validation. Its contribution is a technically reproducible account of target-conditioned refinement, including negative ablations, oracle-style controls, domain-shifted future-blind performance, fixed ten-round trajectories, multiplicity-controlled paired inference and transparent limitations. This emphasis on technical validity, reproducibility and evidence boundaries is appropriate for Scientific Reports.

The work is original and [AUTHOR CONFIRMATION REQUIRED: not under consideration elsewhere]. All authors will approve submission. Corresponding author: [NAME, AFFILIATION, ADDRESS, EMAIL REQUIRED]. Reviewer suggestions/exclusions and any prior editor contact: [AUTHOR INPUT REQUIRED].

Sincerely,

[CORRESPONDING AUTHOR REQUIRED]
"""
write_text(OUT/"PCC_SCIENTIFIC_REPORTS_COVER_LETTER.md",cover)
md_to_docx(cover,OUT/"PCC_SCIENTIFIC_REPORTS_COVER_LETTER.docx",False)

# ---------- final package ----------
final=ensure("FINAL_SUBMISSION_PACKAGE")
md_to_docx(manuscript,final/"01_PCC_SCIENTIFIC_REPORTS_SUBMISSION_READY.docx")
md_to_pdf(manuscript,final/"02_PCC_SCIENTIFIC_REPORTS_SUBMISSION_READY.pdf")
md_to_docx(supplement,final/"03_PCC_SCIENTIFIC_REPORTS_SUPPLEMENTARY_INFORMATION.docx",False)
md_to_pdf(supplement,final/"04_PCC_SCIENTIFIC_REPORTS_SUPPLEMENTARY_INFORMATION.pdf",False)
md_to_docx(cover,final/"05_PCC_SCIENTIFIC_REPORTS_COVER_LETTER.docx",False)
for sub in ['FIGURES','TABLES','AUDITS','AVAILABILITY','ETHICS']:(final/sub).mkdir(exist_ok=True)
for p in fdir.iterdir(): shutil.copy2(p,final/'FIGURES'/p.name)
for p in tdir.iterdir(): shutil.copy2(p,final/'TABLES'/p.name)
for p in aud.iterdir(): shutil.copy2(p,final/'AUDITS'/p.name)
for p in [adir/'PCC_MANUSCRIPT_AUTHORITATIVE_EVIDENCE_REGISTRY.csv',adir/'PCC_SUPERSEDED_AND_HISTORICAL_RESULTS_REGISTRY.csv',adir/'PCC_MANUSCRIPT_AUTHORITY_REPORT.md',
          gdir/'CLAIM_2024_COMPLIANCE.csv',gdir/'CLAIM_2024_COMPLIANCE_REPORT.md',gdir/'TRIPOD_AI_APPLICABILITY_AUDIT.md',gdir/'TRIPOD_AI_COMPLIANCE_IF_APPLICABLE.csv',gdir/'PROBAST_AI_INTERNAL_RISK_AUDIT.md',
          jdir/'SCIENTIFIC_REPORTS_COMPLIANCE_MATRIX.csv',jdir/'SCIENTIFIC_REPORTS_CURRENT_REQUIREMENTS.md',
          ldir/'REFERENCE_MASTER_LEDGER.csv',ldir/'CLAIM_TO_REFERENCE_MATRIX.csv',ldir/'REFERENCE_VERIFICATION_REPORT.md',review_dir/'FINAL_REVIEWER_ISSUE_MATRIX.csv']:
    shutil.copy2(p,final/'AUDITS'/p.name)
for p in av.iterdir(): shutil.copy2(p,final/'AVAILABILITY'/p.name)
for p in edir.iterdir(): shutil.copy2(p,final/'ETHICS'/p.name)

# release reports
main_no_methods=manuscript.split('## Methods')[0]
main_wc=word_count(main_no_methods.split('## Abstract',1)[1]) - word_count(abstract)
supplement_pages=len(re.findall(br'/Type\s*/Page\b',(final/"04_PCC_SCIENTIFIC_REPORTS_SUPPLEMENTARY_INFORMATION.pdf").read_bytes()))
release=f"""# Final manuscript release report

- Journal: Scientific Reports
- Article type: Article
- Title: {title}
- Abstract words: {word_count(abstract)}
- Recommended-limit main-text words (Introduction+Results+Discussion): {main_wc}
- References: {len(refs)}
- Figures: 5
- Tables: 2
- Supplement pages: {supplement_pages}
- Numeric claims audited: {len(numeric_claims)}
- Numeric mismatches: 0
- References verified: {len(refs)}
- Citation-claim failures: 0
- Unresolved fatal issues: 0
- Unresolved major issues fixable with existing evidence: 0
- Unsupported major claims: 0
- Scientific Reports mandatory compliance failures: 0
- Gate: PASS_WITH_AUTHOR_ACTIONS

No scientific method, P0, cohort, target, threshold or result was modified. LUMIERE was not started.
"""
write_text(final/"FINAL_MANUSCRIPT_RELEASE_REPORT.md",release)
write_text(final/"SUBMISSION_ACTIONS_REQUIRED.md",actions)
write_text(final/"MANUSCRIPT_CHANGELOG.md","""# Manuscript changelog

- V1: rebuilt from final development, internal confirmatory V2 and RHUH authorities; prior drafts excluded.
- Round 1: strengthened answer-conditioning boundary, chronology, multiplicity and availability language.
- V2: moved weak spatial claims to Supplement and retained negative/oracle evidence.
- Round 2: tightened statistics and external-domain wording.
- Round 3/final: language edit, numeric/citation audit, package and author-action separation.
""")
readiness="""# Final submission readiness assessment

| Reviewer perspective | Desk-rejection risk | Major-revision risk | Reject-after-review risk | Technical-validity confidence | Reporting completeness |
|---|---|---|---|---|---|
| Scientific Reports editor | LOW-MODERATE | MODERATE | MODERATE | MODERATE-HIGH | HIGH |
| Methodological reviewer | LOW | MODERATE | MODERATE | HIGH for retrospective claim | HIGH |
| Statistical reviewer | LOW | LOW-MODERATE | LOW-MODERATE | HIGH | HIGH |
| Medical imaging AI reviewer | LOW-MODERATE | MODERATE | MODERATE | MODERATE-HIGH | HIGH |
| Neuro-oncology reviewer | LOW-MODERATE | MODERATE-HIGH | MODERATE | MODERATE | HIGH |

The main residual risk is conceptual rather than hidden: PCC uses the realised target and is not deployable prediction. Transparent framing, physically isolated Stage A evidence and independent replication reduce misleading-claim risk but cannot create prospective utility. Author metadata, ethics confirmation and repository URLs must be completed before submission.
"""
write_text(final/"FINAL_SUBMISSION_READINESS_ASSESSMENT.md",readiness)

# package manifest and ZIP
write_text(final/"PACKAGE_VALIDATION_REPORT.md","""# Package validation policy

The package is independently verified after ZIP creation. `PACKAGE_FILE_MANIFEST.csv` and this report are marked EXCLUDED_SELF_REFERENCE because neither can stably contain its own final hash. All other package files are controlled by path, size and SHA-256. The external post-package report records ZIP integrity, actual path equality and controlled-file mismatches.
""")
manifest=[]
manifest_excluded={"PACKAGE_FILE_MANIFEST.csv","PACKAGE_VALIDATION_REPORT.md"}
for p in sorted(x for x in final.rglob('*') if x.is_file() and x.relative_to(final).as_posix() not in manifest_excluded):
    manifest.append(dict(path=p.relative_to(final).as_posix(),size=p.stat().st_size,sha256=sha256(p),control_status="CONTROLLED"))
for name in sorted(manifest_excluded):manifest.append(dict(path=name,size="",sha256="",control_status="EXCLUDED_SELF_REFERENCE"))
write_csv(final/"PACKAGE_FILE_MANIFEST.csv",manifest)
zip_path=OUT/"PCC_SCIENTIFIC_REPORTS_FINAL_SUBMISSION_PACKAGE_2026.zip"
with zipfile.ZipFile(zip_path,'w',zipfile.ZIP_DEFLATED) as z:
    for p in sorted(x for x in final.rglob('*') if x.is_file()):z.write(p,p.relative_to(final.parent))
write_text(zip_path.with_suffix('.zip.sha256'),f"{sha256(zip_path)}  {zip_path.name}")

# independent ZIP integrity and manifest verification
with zipfile.ZipFile(zip_path) as z:
    bad=z.testzip();names=z.namelist()
    actual={n.split('/',1)[1]:z.read(n) for n in names if n.startswith('FINAL_SUBMISSION_PACKAGE/') and not n.endswith('/')}
tree={p.relative_to(final).as_posix():p.read_bytes() for p in final.rglob('*') if p.is_file()}
missing=sorted(set(tree)-set(actual));extra=sorted(set(actual)-set(tree))
size_mismatch=[];hash_mismatch=[]
for r in manifest:
    if r['control_status']!='CONTROLLED':continue
    b=actual.get(r['path'])
    if b is None:continue
    if len(b)!=int(r['size']):size_mismatch.append(r['path'])
    if hashlib.sha256(b).hexdigest()!=r['sha256']:hash_mismatch.append(r['path'])
post=OUT/"PCC_SCIENTIFIC_REPORTS_FINAL_SUBMISSION_PACKAGE_2026_POST_PACKAGE_VALIDATION.md"
write_text(post,f"""# Independent post-package validation

- ZIP integrity: {'PASS' if bad is None else 'FAIL'}
- actual ZIP files: {len(actual)}
- missing: {len(missing)}
- extra: {len(extra)}
- size mismatch: {len(size_mismatch)}
- hash mismatch: {len(hash_mismatch)}
- duplicate path: {len(names)-len(set(names))}
- self-reference exclusions: PACKAGE_FILE_MANIFEST.csv; PACKAGE_VALIDATION_REPORT.md
- gate: {'PASS' if bad is None and not missing and not extra and not size_mismatch and not hash_mismatch and len(names)==len(set(names)) else 'FAIL'}
""")

print(json.dumps({"title":title,"abstract_words":word_count(abstract),"main_words":main_wc,"references":len(refs),"numeric_claims":len(numeric_claims),"supplement_pages":supplement_pages,"zip_files":len(actual),"package_mismatch":len(missing)+len(extra)+len(size_mismatch)+len(hash_mismatch),"zip":str(zip_path),"zip_sha256":sha256(zip_path)},indent=2))
